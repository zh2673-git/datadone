#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Word导出器 - 生成解读式初核分析报告

报告结构（5节 + 附录，解读而非堆砌）：
  一、数据概况 — 数据覆盖与可靠性评估
  二、资金全貌 — 各平台收支解读、存取现解读（引用Excel存取现明细表）、关键收支特征含特殊金额/日期（引用Excel重点收支明细表）
  三、重点发现 — 核心交易对手（表格）+ 行为发现简述（异常/模式/大额流向合为一段）
  四、风险研判与调查建议 — 结论导向
  附录：术语说明（含交叉类型、行为发现、规避阈值、特殊金额等新术语）

设计原则：
  - 解读而非报数：不说"存现XX元"，而说"存取现占收支比X%，需关注"
  - 聚合而非罗列：异常/模式/大额流向合为简述段落，引导查阅Excel详表
  - 指向而非复述：Word聚焦初核关注点和关键数据，详表均在Excel中
"""

import logging
import os
from typing import Optional
from datetime import datetime
from collections import Counter

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.models.report import ReportData, ExportConfig, PersonReportData
from src.models.analysis_result import (
    PersonBaseline, AnomalyItem, PatternMatch, TimelineChain, RiskAssessment
)

logger = logging.getLogger(__name__)


# ================================================================== #
#  术语说明表
# ================================================================== #

GLOSSARY = [
    ("行为基线", "根据历史交易数据统计的月均收支水平、交易频次、对手数量等参考值，作为判断异常的对照基准。在Excel\"风险研判表\"前半部分呈现"),
    ("σ（标准差）", "衡量数据波动幅度的统计单位。偏离基线越多个σ，说明该行为越异常。一般≥2σ即值得关注，≥3σ为显著异常"),
    ("偏离度", "某项指标偏离行为基线的程度，以σ倍数表示。如偏离3.2σ表示超出正常波动范围3.2个标准差"),
    ("P25/P75", "第25/75百分位数。P75常作为判断'大额'的参考阈值——超过75%的历史交易即为偏高"),
    ("匹配度", "行为模式的匹配程度，0-100%。越高说明该行为越符合某种可疑模式"),
    ("周期性收入", "与同一对手在连续多个月内定期收到金额相近的资金，常见于工资、利益输送等场景"),
    ("资金中转", "收到资金后短时间内（48小时内）等额或近似等额转出，资金未沉淀，疑似过账"),
    ("规避检测", "将大额交易拆分为多笔略低于监管阈值的交易（如4-5万），以规避大额交易报告制度"),
    ("时序链", "通话与资金往来在时间上紧密关联的序列，如'通话→收入→通话'，是利益输送的强证据"),
    ("纯收入/纯支出对手", "只存在单向资金往来的交易对手——要么只收不付，要么只付不收"),
    ("存取现", "银行现金存款（存现）和现金取款（取现）。大额频繁存取现可能涉及规避资金追踪。详见Excel\"存取现明细表\""),
    ("同日存取", "同一日期既有存现又有取现操作，可能涉及资金洗转或规避监管"),
    ("连续取现", "连续多日取现，可能涉及分散提取大额现金以规避单笔大额交易报告"),
    ("净流向", "收入减去支出的差额。正值为净流入，负值为净流出"),
    ("严重程度", "异常行为的严重性评级。'高'表示偏离度≥3σ或涉及大额资金，'中'为2-3σ，'低'为其他"),
    ("证据充分度", "当前数据对调查结论的支撑程度。'充分'可形成完整证据链，'较充分'需少量补充，'待补充'有明显缺口"),
    ("大额交易报告", "金融机构对单笔5万元以上现金交易或50万元以上转账交易向反洗钱中心报送的制度"),
    ("交叉类型", "综合分析表中对手的跨平台关联程度。'双平台交叉'表示既有通话又有资金往来，初核价值最高"),
    ("行为发现", "行为模式识别与时序链证据的统称，在Excel\"行为发现表\"中统一呈现"),
    ("规避阈值", "交易金额接近但未达到监管报告阈值的交易（如4-5万接近5万现金报告线、40-50万接近50万转账报告线），可能暗示有意识规避监管。详见Excel\"重点收支明细表\""),
    ("特殊金额交易", "包含特殊数字特征的交易金额，如520/1314爱情数字（暗示非正常商业关系）、8888/9999吉利数字（常见于利益输送）。详见Excel\"重点收支明细表\""),
    ("特殊日期交易", "节假日期间的大额交易，可能是\"礼尚往来\"形式下的利益输送。详见Excel\"重点收支明细表\""),
    ("习惯等级", "根据消费频次判定的行为稳定性。'高频习惯'（月均≥4次）为稳定生活规律，'低频偏好'（月均1-3次）为有意识消费选择，'偶尔消费'（月均<1次）不足以判定为习惯。详见Excel\"习惯兴趣表\""),
    ("典型时段", "某类消费最常发生的时间段。如'深夜'表示该类消费多在22点后发生，可能反映作息特征。详见Excel\"习惯兴趣表\""),
    ("商户识别", "通过对方姓名中的商户名称特征（如'XX健身房'）识别消费类别，与关键词匹配互为补充。详见Excel\"习惯兴趣表\""),
]


class WordExporter:
    """Word报告导出器 - 解读式初核分析报告"""

    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)

    def export(self, report_data: ReportData, config: Optional[ExportConfig] = None) -> str:
        if config is None:
            config = ExportConfig()

        output_dir = config.output_dir
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, "分析报告.docx")

        self.logger.info("开始生成Word报告...")

        doc = Document()

        # 标题页
        self._write_title_page(doc, report_data)

        # 一、数据概况
        self._write_data_overview(doc, report_data)

        # 二~四：逐人分析
        for idx, person in enumerate(report_data.persons, 1):
            person_data = report_data.person_reports.get(person)
            if person_data is None:
                continue
            doc.add_heading(f"{idx}、{person}", level=2)
            self._write_financial_overview(doc, person_data, report_data)
            self._write_key_findings(doc, person_data, report_data)
            self._write_ai_insights(doc, person_data)
            self._write_risk_and_suggestions(doc, person_data)

        # 附录：术语说明
        self._write_glossary(doc)

        doc.save(output_path)
        self.logger.info(f"Word报告已生成: {output_path}")
        return output_path

    # ================================================================== #
    #  标题页
    # ================================================================== #

    def _write_title_page(self, doc: Document, report_data: ReportData) -> None:
        person_names = "、".join(report_data.persons) if report_data.persons else "未知"
        title = f"{person_names}线索技术分析报告"
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"分析时间：{datetime.now().strftime('%Y年%m月%d日')}")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("（本报告为数据分析技术辅助结果，供调查参考，不作为结论性依据）")
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ================================================================== #
    #  一、数据概况
    # ================================================================== #

    def _write_data_overview(self, doc: Document, report_data: ReportData) -> None:
        doc.add_heading("一、数据概况", level=2)

        # 被分析人
        p = doc.add_paragraph()
        run = p.add_run("被分析人：")
        run.bold = True
        p.add_run("、".join(report_data.persons) if report_data.persons else "无")

        # 数据类型
        type_map = {'bank': '银行', 'wechat': '微信', 'alipay': '支付宝', 'call': '话单'}
        data_type_names = [type_map.get(dt, dt) for dt in report_data.data_types]
        p = doc.add_paragraph()
        run = p.add_run("数据类型：")
        run.bold = True
        p.add_run("、".join(data_type_names) if data_type_names else "无")

        # 涉及银行
        if report_data.bank_names:
            p = doc.add_paragraph()
            run = p.add_run("涉及银行：")
            run.bold = True
            p.add_run("、".join(report_data.bank_names))

        # 数据可靠性评估（基于基线）
        for person in report_data.persons:
            pr = report_data.person_reports.get(person)
            if pr and pr.行为基线:
                bl = pr.行为基线
                p = doc.add_paragraph()
                run = p.add_run(f"{person}数据可靠性：")
                run.bold = True

                if bl.数据充足度 == '充足':
                    p.add_run(f"数据充足（覆盖{bl.数据月数}个月），基线参考性较强")
                elif bl.数据充足度 == '一般':
                    p.add_run(f"数据一般（覆盖{bl.数据月数}个月），基线仅供参考，结论需谨慎")
                else:
                    p.add_run(f"数据不足（仅覆盖{bl.数据月数}个月），基线参考性弱，需补充数据")

    # ================================================================== #
    #  二、资金全貌（逐人）
    # ================================================================== #

    def _write_financial_overview(self, doc: Document, person: PersonReportData,
                                   report_data: ReportData) -> None:
        """资金全貌 — 解读式呈现，非报数"""
        doc.add_heading("二、资金全貌", level=3)

        # ---- 2.1 收支总览 ----
        doc.add_heading("收支总览", level=4)

        total_income = person.银行总进账 + person.微信总进账 + person.支付宝总进账
        total_expense = person.银行总出账 + person.微信总出账 + person.支付宝总出账
        net_flow = total_income - total_expense

        # 汇总表
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ['平台', '收入', '支出', '净流向']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for par in cell.paragraphs:
                for run in par.runs:
                    run.bold = True

        platforms = []
        if person.银行总进账 > 0 or person.银行总出账 > 0:
            platforms.append(('银行', person.银行总进账, person.银行总出账))
        if person.微信总进账 > 0 or person.微信总出账 > 0:
            platforms.append(('微信', person.微信总进账, person.微信总出账))
        if person.支付宝总进账 > 0 or person.支付宝总出账 > 0:
            platforms.append(('支付宝', person.支付宝总进账, person.支付宝总出账))

        for name, inc, exp in platforms:
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = f"{inc:,.0f}"
            row_cells[2].text = f"{exp:,.0f}"
            net = inc - exp
            row_cells[3].text = f"{'净流入' if net >= 0 else '净流出'}{abs(net):,.0f}"

        # 合计行
        row_cells = table.add_row().cells
        for par in row_cells[0].paragraphs:
            for run in par.runs:
                run.bold = True
        row_cells[0].text = "合计"
        row_cells[1].text = f"{total_income:,.0f}"
        row_cells[2].text = f"{total_expense:,.0f}"
        row_cells[3].text = f"{'净流入' if net_flow >= 0 else '净流出'}{abs(net_flow):,.0f}"

        doc.add_paragraph()  # 表后空行

        # 解读文字
        bl = person.行为基线
        p = doc.add_paragraph()
        if bl and bl.数据月数 > 0:
            p.add_run(
                f"数据跨度{person.时间跨度}（{bl.数据月数}个月），"
                f"月均收入{bl.月均收入:,.0f}元、支出{bl.月均支出:,.0f}元。"
            )
            # 收支平衡解读
            if net_flow > 0:
                p.add_run(f"总体净流入{net_flow:,.0f}元")
                if bl.月均收入 > 0:
                    months = bl.数据月数
                    monthly_net = net_flow / months
                    if monthly_net > bl.月均收入 * 0.5:
                        p.add_run("，净流入占收入比例较高，需关注资金来源")
                    else:
                        p.add_run("，收支基本平衡")
            else:
                p.add_run(f"总体净流出{abs(net_flow):,.0f}元")
                if abs(net_flow) > total_income * 0.3:
                    p.add_run("，支出显著大于收入，需关注资金来源补充渠道")

            # 趋势
            trends = []
            if bl.收入趋势 != '平稳':
                trends.append(f"收入呈{bl.收入趋势}趋势")
            if bl.支出趋势 != '平稳':
                trends.append(f"支出呈{bl.支出趋势}趋势")
            if trends:
                p.add_run(f"。{('，'.join(trends))}")

        # ---- 2.2 存取现 ----
        self._write_cash_overview(doc, person)

        # ---- 2.3 关键收支特征 ----
        self._write_key_income_overview(doc, person)

    def _write_cash_overview(self, doc: Document, person: PersonReportData) -> None:
        """存取现解读 — 初核重点关注"""
        if person.存现总额 == 0 and person.取现总额 == 0:
            return

        doc.add_heading("存取现情况", level=4)

        bl = person.行为基线
        total_flow = person.银行总进账 + person.银行总出账

        # 基本数据
        p = doc.add_paragraph()
        p.add_run(f"存现{person.存现总额:,.0f}元，取现{person.取现总额:,.0f}元")

        # 比例解读
        if bl and bl.存取现占收支比 > 0:
            p.add_run(f"，存取现占收支比{bl.存取现占收支比:.1f}%")
            if bl.存取现占收支比 > 20:
                p.add_run("（偏高，需关注是否有规避转账追踪的意图）")
            elif bl.存取现占收支比 > 10:
                p.add_run("（中等水平）")
            else:
                p.add_run("（正常水平）")

        # 大额存取现
        large_cash = person.单笔万以上存取现
        if large_cash and large_cash.get('次数', 0) > 0:
            p.add_run(
                f"。其中单笔万元以上存取现{large_cash['次数']}次、"
                f"合计{large_cash['金额']:,.0f}元"
            )
            if large_cash['金额'] > total_flow * 0.1:
                p.add_run("，大额现金操作占比较高")

        # 存取现话单匹配
        if person.存取现话单匹配:
            p = doc.add_paragraph()
            match_count = len(person.存取现话单匹配)
            # 统计有通话对方的匹配
            with_counterparty = [m for m in person.存取现话单匹配 if m.get('通话对方')]
            p.add_run(f"存取现前后有通话记录{match_count}次")
            if with_counterparty:
                p.add_run(f"，其中{len(with_counterparty)}次可关联到通话对方")
                # 展示前3条关键匹配
                for m in with_counterparty[:3]:
                    detail = doc.add_paragraph()
                    detail.paragraph_format.left_indent = Inches(0.3)
                    desc = f"{m.get('交易日期', '')} {m.get('交易方向', '')}{m.get('交易金额', 0):,.0f}元 → 通话对方：{m.get('通话对方', '未知')}"
                    if m.get('通话对方单位'):
                        desc += f"（{m['通话对方单位']}）"
                    detail.add_run(desc)

        # 提示：详见Excel存取现明细表
        p = doc.add_paragraph()
        p.add_run('（逐笔存取现明细、大额标记等见Excel"存取现明细表"）')
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    def _write_key_income_overview(self, doc: Document, person: PersonReportData) -> None:
        """关键收支特征概述 — 初核重点关注"""
        items = []
        if person.工作收入:
            kw = person.工作收入.get('匹配关键词', '')
            total = person.工作收入.get('总额', 0)
            count = person.工作收入.get('笔数', 0)
            items.append(f"工作收入{total:,.0f}元" + (f"（{count}笔，疑似{kw}）" if kw and count else f"（{count}笔）" if count else ""))
        if person.房产收入 and person.房产收入.get('总额', 0) > 0:
            items.append(f"房产收入{person.房产收入.get('总额', 0):,.0f}元")
        if person.房产支出 and person.房产支出.get('总额', 0) > 0:
            items.append(f"房产支出{person.房产支出.get('总额', 0):,.0f}元")
        if person.租金收入 and person.租金收入.get('总额', 0) > 0:
            count = person.租金收入.get('笔数', 0)
            items.append(f"租金收入{person.租金收入.get('总额', 0):,.0f}元" + (f"（{count}笔）" if count else ""))
        if person.租金支出 and person.租金支出.get('总额', 0) > 0:
            items.append(f"租金支出{person.租金支出.get('总额', 0):,.0f}元")
        if person.车辆收入 and person.车辆收入.get('总额', 0) > 0:
            items.append(f"车辆收入{person.车辆收入.get('总额', 0):,.0f}元")
        if person.车辆支出 and person.车辆支出.get('总额', 0) > 0:
            items.append(f"车辆支出{person.车辆支出.get('总额', 0):,.0f}元")
        if person.证券收入 and person.证券收入.get('总额', 0) > 0:
            items.append(f"证券收入{person.证券收入.get('总额', 0):,.0f}元")
        if person.证券支出 and person.证券支出.get('总额', 0) > 0:
            items.append(f"证券支出{person.证券支出.get('总额', 0):,.0f}元")

        if not items:
            return

        doc.add_heading("关键收支特征", level=4)

        # 主要收支概述
        p = doc.add_paragraph()
        p.add_run("；".join(items))

        # 对比月均收入 — 收入来源差异
        bl = person.行为基线
        if bl and bl.月均收入 > 0 and person.工作收入:
            work_annual = person.工作收入.get('总额', 0)
            if bl.数据月数 > 0:
                work_monthly = work_annual / bl.数据月数
                if work_monthly < bl.月均收入 * 0.3:
                    p2 = doc.add_paragraph()
                    run = p2.add_run("⚠ 收入来源差异：")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0, 0)
                    p2.add_run(
                        f"可识别工作收入月均约{work_monthly:,.0f}元，"
                        f"远低于实际月均收入{bl.月均收入:,.0f}元，"
                        f"存在{(bl.月均收入 - work_monthly):,.0f}元/月的收入来源需核实"
                    )

        # 特殊金额交易摘要
        love_trades = person.爱情数字交易
        other_special = person.其他特殊金额交易
        if love_trades or other_special:
            p = doc.add_paragraph()
            run = p.add_run("特殊金额交易：")
            run.bold = True
            summaries = []
            if love_trades:
                total_amt = sum(t.get('金额', 0) for t in love_trades)
                summaries.append(f"爱情数字{len(love_trades)}笔共{total_amt:,.0f}元（如520/1314等，暗示非正常商业关系）")
            if other_special:
                total_amt = sum(t.get('金额', 0) for t in other_special)
                summaries.append(f"其他特殊金额{len(other_special)}笔共{total_amt:,.0f}元（如8888/9999等，常见于利益输送）")
            p.add_run("；".join(summaries))

        # 特殊日期交易摘要
        special_dates = person.特殊日期交易
        if special_dates:
            p = doc.add_paragraph()
            run = p.add_run("特殊日期交易：")
            run.bold = True
            date_counter = Counter()
            total_amt = 0
            for t in special_dates:
                name = t.get('日期名称', '特殊日期')
                date_counter[name] += 1
                total_amt += t.get('金额', 0)
            date_summary = "、".join(f"{name}{cnt}笔" for name, cnt in date_counter.most_common(5))
            p.add_run(f"节假日期间大额交易{len(special_dates)}笔共{total_amt:,.0f}元（{date_summary}）")

        # 提示：详见Excel重点收支明细表
        p = doc.add_paragraph()
        p.add_run('（逐笔明细、规避阈值交易、特殊金额/日期交易等见Excel"重点收支明细表"）')
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ================================================================== #
    #  三、重点发现（逐人）
    # ================================================================== #

    def _write_key_findings(self, doc: Document, person: PersonReportData,
                             report_data: ReportData) -> None:
        """重点发现 — 聚合呈现，非逐条罗列"""
        doc.add_heading("三、重点发现", level=3)

        # 3.1 核心交易对手
        self._write_counterparty_overview(doc, person, report_data)

        # 3.2 行为发现简述（异常+模式+大额流向 合并为一段）
        self._write_behavior_brief(doc, person)

        # 3.3 习惯兴趣画像
        self._write_habit_interest_brief(doc, person)

    def _write_counterparty_overview(self, doc: Document, person: PersonReportData,
                                      report_data: ReportData) -> None:
        """核心交易对手概述 — 从频率表和综合分析表提炼"""
        doc.add_heading("核心交易对手", level=4)

        # 从频率表提取该人的TOP5对手
        top_counterparties = []
        for item in report_data.bill_frequency:
            if item.get('本方姓名') == person.本方姓名:
                total = item.get('收入总额', 0) + item.get('支出总额', 0)
                top_counterparties.append({
                    '姓名': item.get('对方姓名', ''),
                    '收入': item.get('收入总额', 0),
                    '支出': item.get('支出总额', 0),
                    '总金额': total,
                    '平台': item.get('平台', ''),
                })

        # 去重+合并
        merged = {}
        for c in top_counterparties:
            name = c['姓名']
            if name in merged:
                merged[name]['收入'] += c['收入']
                merged[name]['支出'] += c['支出']
                merged[name]['总金额'] += c['总金额']
                if c['平台'] not in merged[name]['平台']:
                    merged[name]['平台'] += f"、{c['平台']}"
            else:
                merged[name] = dict(c)

        top_list = sorted(merged.values(), key=lambda x: x['总金额'], reverse=True)[:10]

        if top_list:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, h in enumerate(['对手姓名', '收入总额', '支出总额', '涉及平台', '备注']):
                cell = table.rows[0].cells[i]
                cell.text = h
                for par in cell.paragraphs:
                    for run in par.runs:
                        run.bold = True

            # 查综合分析表补充对方单位信息
            counterparty_units = {}
            for base_name, items in report_data.cross_analysis.items():
                for item in items:
                    item_dict = item if isinstance(item, dict) else item.model_dump()
                    if item_dict.get('本方姓名') == person.本方姓名:
                        name = item_dict.get('对方姓名', '')
                        unit = item_dict.get('对方单位名称', '') or item_dict.get('对方单位', '')
                        title = item_dict.get('对方职务', '')
                        if name and (unit or title):
                            counterparty_units[name] = f"{unit} {title}".strip()

            for c in top_list:
                row_cells = table.add_row().cells
                row_cells[0].text = str(c['姓名'])
                row_cells[1].text = f"{c['收入']:,.0f}"
                row_cells[2].text = f"{c['支出']:,.0f}"
                row_cells[3].text = str(c['平台'])
                note = counterparty_units.get(c['姓名'], '')
                # 单方向标记
                if c['收入'] > 0 and c['支出'] == 0:
                    note += "；纯收方"
                elif c['支出'] > 0 and c['收入'] == 0:
                    note += "；纯付方"
                row_cells[4].text = note

            doc.add_paragraph()

        # 纯收支对手提示
        pure_income = person.纯收入对手[:3]
        pure_expense = person.纯支出对手[:3]
        if pure_income or pure_expense:
            p = doc.add_paragraph()
            run = p.add_run("单方向资金对手：")
            run.bold = True
            if pure_income:
                names = "、".join([f"{o['对方姓名']}（{o['金额']:,.0f}元）" for o in pure_income])
                p.add_run(f"仅收方—{names}")
                if pure_expense:
                    p.add_run("；")
            if pure_expense:
                names = "、".join([f"{o['对方姓名']}（{o['金额']:,.0f}元）" for o in pure_expense])
                p.add_run(f"仅付方—{names}")

    def _write_behavior_brief(self, doc: Document, person: PersonReportData) -> None:
        """行为发现简述 — 异常+模式+大额流向合为一段，引导查阅Excel"""
        doc.add_heading("行为发现简述", level=4)

        parts = []

        # --- 异常行为 ---
        anomalies = person.异常列表
        if anomalies:
            type_counter = Counter()
            high_count = 0
            for a in anomalies:
                type_counter[a.异常类型] += 1
                if a.严重程度 == '高':
                    high_count += 1
            type_summary = "、".join(
                f"{t}{c}项" for t, c in type_counter.most_common(4)
            )
            anomaly_text = f"检测到异常行为{len(anomalies)}项（{type_summary}）"
            if high_count > 0:
                anomaly_text += f"，其中高严重{high_count}项"
            parts.append(anomaly_text)

        # --- 行为模式 ---
        patterns = person.行为模式
        if patterns:
            pattern_counter = Counter()
            for pat in patterns:
                pattern_counter[pat.模式名称] += 1
            max_match = max(p.匹配度 for p in patterns)
            pattern_summary = "、".join(
                f"{name}{cnt}次" for name, cnt in pattern_counter.most_common(4)
            )
            pattern_text = f"识别行为模式{len(patterns)}次（{pattern_summary}，最高匹配度{max_match:.0%}）"
            parts.append(pattern_text)

        # --- 大额资金流向 ---
        levels = person.大额资金跟踪层级
        if levels:
            level_counter = Counter()
            for item in levels:
                lv = item.get('追踪层级', 0)
                level_counter[lv] += 1
            level_desc = {0: '直接', 1: '一层间接', 2: '二层间接'}
            level_summary = "、".join(
                f"{level_desc.get(lv, f'层级{lv}')}{cnt}条" for lv, cnt in sorted(level_counter.items())
            )
            parts.append(f"大额资金跟踪{len(levels)}条（{level_summary}）")

        if not parts:
            doc.add_paragraph("未检测到显著异常行为或可疑模式。")
            return

        # 合并为一段
        p = doc.add_paragraph()
        p.add_run("；".join(parts) + "。")

        # 高严重异常：只列最关键的3条
        if anomalies:
            high_anomalies = [a for a in anomalies if a.严重程度 == '高']
            if high_anomalies:
                for a in high_anomalies[:3]:
                    detail = doc.add_paragraph()
                    detail.paragraph_format.left_indent = Inches(0.3)
                    run = detail.add_run(f"⚠ [{a.异常子类}] {a.说明}")
                    run.font.color.rgb = RGBColor(0xFF, 0, 0)
                    if a.对方姓名:
                        detail.add_run(f"（对方：{a.对方姓名}）")

        # 行为模式：只列最高匹配度的1-2条关键证据
        if patterns:
            best_patterns = sorted(patterns, key=lambda p: p.匹配度, reverse=True)[:2]
            for pat in best_patterns:
                detail = doc.add_paragraph()
                detail.paragraph_format.left_indent = Inches(0.3)
                run = detail.add_run(f"◆ {pat.模式名称}（匹配度{pat.匹配度:.0%}）")
                run.bold = True
                if pat.涉及对手:
                    detail.add_run(f" — 涉及{pat.涉及对手}")

        # 引导查阅Excel
        p = doc.add_paragraph()
        p.add_run('（详细异常明细、行为模式匹配与时序链证据见Excel"异常明细表"和"行为发现表"，大额资金逐笔跟踪见"大额资金跟踪表"）')
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ================================================================== #
    #  3.3 习惯兴趣画像
    # ================================================================== #

    def _write_habit_interest_brief(self, doc: Document, person: PersonReportData) -> None:
        """习惯兴趣画像 — 按习惯等级分层呈现"""
        habits = person.习惯兴趣
        if not habits:
            return

        doc.add_heading("习惯兴趣画像", level=4)

        high_freq = [h for h in habits if h.习惯等级 == "高频习惯"]
        low_freq = [h for h in habits if h.习惯等级 == "低频偏好"]
        occasional = [h for h in habits if h.习惯等级 == "偶尔消费"]

        if high_freq:
            p = doc.add_paragraph()
            run = p.add_run("稳定习惯：")
            run.bold = True
            items = []
            for h in high_freq[:8]:
                desc = f"{h.类别}（{h.子类}，月均{h.月均频次:.1f}次"
                if h.典型时段:
                    desc += f"，{h.典型时段}"
                desc += "）"
                items.append(desc)
            p.add_run("、".join(items))

        if low_freq:
            p = doc.add_paragraph()
            run = p.add_run("消费偏好：")
            run.bold = True
            items = []
            for h in low_freq[:6]:
                desc = f"{h.类别}（{h.子类}，月均{h.月均频次:.1f}次"
                if h.总金额 > 0:
                    desc += f"，累计{h.总金额:,.0f}元"
                desc += "）"
                items.append(desc)
            p.add_run("、".join(items))

        if occasional:
            p = doc.add_paragraph()
            run = p.add_run("偶尔消费：")
            run.bold = True
            items = []
            for h in occasional[:4]:
                items.append(f"{h.类别}（{h.子类}）")
            p.add_run("、".join(items))

        p = doc.add_paragraph()
        p.add_run('（习惯兴趣详细数据见Excel"习惯兴趣表"）')
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ================================================================== #
    #  四、风险研判与调查建议（逐人）
    # ================================================================== #

    def _write_risk_and_suggestions(self, doc: Document, person: PersonReportData) -> None:
        """风险研判与调查建议"""
        doc.add_heading("四、风险研判与调查建议", level=3)

        ra = person.风险研判
        if ra is None:
            doc.add_paragraph("风险研判数据不足。")
            return

        # 4.1 风险等级
        p = doc.add_paragraph()
        run = p.add_run("综合风险等级：")
        run.bold = True

        level_color = {
            "高风险": RGBColor(0xFF, 0, 0),
            "中风险": RGBColor(0xFF, 0x8C, 0),
            "低风险": RGBColor(0, 0x80, 0),
        }
        level_run = p.add_run(f"{ra.综合风险等级}（{ra.综合风险分数:.1f}/100分）")
        level_run.bold = True
        level_run.font.color.rgb = level_color.get(ra.综合风险等级, RGBColor(0, 0, 0))

        # 四维度评分
        p = doc.add_paragraph()
        p.add_run("评分构成：")
        dims = []
        if ra.异常偏离度得分 > 0:
            dims.append(f"异常偏离{ra.异常偏离度得分:.0f}分")
        if ra.行为模式得分 > 0:
            dims.append(f"行为模式{ra.行为模式得分:.0f}分")
        if ra.证据链得分 > 0:
            dims.append(f"证据链{ra.证据链得分:.0f}分")
        if ra.规模得分 > 0:
            dims.append(f"涉及规模{ra.规模得分:.0f}分")
        p.add_run(" + ".join(dims))

        # 4.2 证据充分度
        p = doc.add_paragraph()
        run = p.add_run("证据充分度：")
        run.bold = True

        sufficiency_color = {
            "充分": RGBColor(0, 0x80, 0),
            "较充分": RGBColor(0xFF, 0x8C, 0),
            "待补充": RGBColor(0xFF, 0, 0),
        }
        suff_run = p.add_run(ra.证据充分度)
        suff_run.font.color.rgb = sufficiency_color.get(ra.证据充分度, RGBColor(0, 0, 0))

        if ra.已有证据:
            detail = doc.add_paragraph()
            detail.paragraph_format.left_indent = Inches(0.3)
            detail.add_run("已有证据：" + "；".join(ra.已有证据))

        if ra.待补充证据:
            detail = doc.add_paragraph()
            detail.paragraph_format.left_indent = Inches(0.3)
            detail.add_run("待补充证据：" + "；".join(ra.待补充证据))

        # 4.3 重点人员
        if ra.重点人员:
            doc.add_heading("重点人员", level=4)
            for kp in ra.重点人员:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                unit_info = ""
                if kp.对方单位:
                    unit_info = f"（{kp.对方单位}"
                    if kp.对方职务:
                        unit_info += f" {kp.对方职务}"
                    unit_info += "）"
                run = p.add_run(f"TOP{kp.排名} {kp.对方姓名}{unit_info}")
                run.bold = True

                details = []
                if kp.通话次数 > 0:
                    details.append(f"通话{kp.通话次数}次")
                if kp.资金往来金额 > 0:
                    details.append(f"资金往来{kp.资金往来金额:,.0f}元")
                if kp.时序链数 > 0:
                    details.append(f"{kp.时序链数}条时序链")
                if kp.关联特征:
                    details.append(kp.关联特征)
                if details:
                    p.add_run(" — " + "；".join(details))

        # 4.4 调查方向
        if ra.调查方向建议:
            doc.add_heading("调查方向", level=4)
            for idx, suggestion in enumerate(ra.调查方向建议, 1):
                detail = doc.add_paragraph()
                detail.paragraph_format.left_indent = Inches(0.3)
                detail.add_run(f"({idx}) {suggestion}")

        # 4.5 重点时段
        if ra.重点时段:
            p = doc.add_paragraph()
            run = p.add_run("重点时段：")
            run.bold = True
            p.add_run(ra.重点时段)

    def _write_ai_insights(self, doc: Document, person: PersonReportData) -> None:
        """AI 智能解读（基于规则信号的叙事与假设）"""
        if not person.ai_insights:
            return

        insight = person.ai_insights
        has_content = (
            insight.人员级叙事 or
            insight.线索级叙事 or
            insight.调查假设
        )
        if not has_content:
            return

        doc.add_heading("（AI）智能解读与调查提示", level=3)

        # 人员级叙事
        if insight.人员级叙事:
            doc.add_heading("AI 综合叙事", level=4)
            for block in insight.人员级叙事:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(f"{block.标题}（置信度：{block.置信度}）")
                run.bold = True
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.add_run(block.内容)

        # 调查假设
        if insight.调查假设:
            doc.add_heading("AI 调查假设", level=4)
            for idx, hypo in enumerate(insight.调查假设, 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(f"假设 {idx}：{hypo.标题}（风险等级：{hypo.风险等级}）")
                run.bold = True
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.add_run(hypo.描述)
                if hypo.验证方向:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.add_run("建议核实方向：" + "；".join(hypo.验证方向))

        # 证据来源提示
        p = doc.add_paragraph()
        p.add_run("【说明】以上内容由 AI 根据规则分析结果生成，仅供调查参考，"
                  "具体结论需结合原始数据人工复核。").italic = True

    # ================================================================== #
    #  附录：术语说明
    # ================================================================== #

    def _write_glossary(self, doc: Document) -> None:
        doc.add_heading("附录：术语说明", level=2)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for i, h in enumerate(['术语', '说明']):
            cell = table.rows[0].cells[i]
            cell.text = h
            for par in cell.paragraphs:
                for run in par.runs:
                    run.bold = True

        for term, explanation in GLOSSARY:
            row_cells = table.add_row().cells
            row_cells[0].text = term
            row_cells[1].text = explanation
            # 术语列加粗
            for par in row_cells[0].paragraphs:
                for run in par.runs:
                    run.bold = True
