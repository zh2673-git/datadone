#!/usr/bin/env python
# -*- coding: utf-8 -*-

import pandas as pd
import os
import logging
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np
from collections import defaultdict

from src.utils.fund_tracking import FundTrackingEngine
from src.utils.key_transactions import KeyTransactionEngine
from src.utils.config import Config

class NewWordExporter:
    def __init__(self, output_dir: str = 'output', config=None):
        self.logger = logging.getLogger(self.__class__.__name__)
        self.output_dir = output_dir
        self.config = config or Config()
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def generate_comprehensive_report(self, report_title: str, data_models: Dict, analyzers: Dict):
        """
        生成新的统一的、以人为核心的综合分析报告。
        """
        print("开始生成新版Word报告...")
        total_steps = 6
        current_step = 0
        
        def update_progress(step_name):
            nonlocal current_step
            current_step += 1
            progress = (current_step / total_steps) * 100
            print(f"[{current_step}/{total_steps}] {step_name}... ({progress:.1f}%)")
        
        doc = Document()
        doc.add_heading(report_title, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 分析数据类型
        data_types = []
        if data_models.get('bank') and not data_models['bank'].data.empty:
            data_types.append('银行')
        if data_models.get('wechat') and not data_models['wechat'].data.empty:
            data_types.append('微信')
        if data_models.get('alipay') and not data_models['alipay'].data.empty:
            data_types.append('支付宝')
        if data_models.get('call') and not data_models['call'].data.empty:
            data_types.append('话单')
            
        # 获取所有涉及的人员
        update_progress("识别分析对象")
        persons_with_financials = self._get_persons_with_financial_data(data_models)
        
        if not persons_with_financials:
            doc.add_paragraph("在所有数据中未能识别出任何持有金融账户（银行、微信、支付宝）的分析对象，无法生成个人详细报告。")
            self._save_document(doc, report_title)
            return

        # 为每个人员生成分析报告
        update_progress("生成个人分析报告")
        for i, person_name in enumerate(persons_with_financials):
            print(f"  正在分析 {person_name} ({i+1}/{len(persons_with_financials)})")
            doc.add_heading(f"{i+1}、{person_name}分析", level=2)
            
            # 在每个被分析人名下添加数据类型说明
            if data_types:
                doc.add_paragraph(f"分析依据的数据类型：{', '.join(data_types)}")
                
            # 如果有银行数据，列出具体银行（针对该人员）
            if '银行' in data_types:
                bank_data = data_models['bank'].data[data_models['bank'].data[data_models['bank'].name_column] == person_name]
                if not bank_data.empty and '银行类型' in bank_data.columns:
                    bank_names = set()
                    bank_names.update(bank_data['银行类型'].dropna().unique())
                    if bank_names:
                        doc.add_paragraph(f"涉及银行：{', '.join(sorted(list(bank_names)))}")
            
            # 一、资金体量分析
            print(f"    1/4 资金体量分析...")
            self._generate_fund_volume_analysis(doc, person_name, data_models, analyzers)
            
            # 二、特殊资金分析
            print(f"    2/4 特殊资金分析...")
            self._generate_special_fund_analysis(doc, person_name, data_models, analyzers)
            
            # 三、重点收支分析
            print(f"    3/4 重点收支分析...")
            self._generate_key_transactions_analysis(doc, person_name, data_models, analyzers)
            
            # 四、重点人员分析
            print(f"    4/4 重点人员分析...")
            self._generate_key_persons_analysis(doc, data_models, analyzers)

        update_progress("保存Word文档")
        self._save_document(doc, report_title)

    def _get_persons_with_financial_data(self, data_models: Dict) -> List[str]:
        """获取所有金融数据源中的所有不重复的本方姓名"""
        financial_persons = set()
        financial_data_types = ['bank', 'wechat', 'alipay']
        for data_type in financial_data_types:
            model = data_models.get(data_type)
            if model and not model.data.empty and hasattr(model, 'name_column'):
                name_col = model.name_column
                if name_col in model.data.columns:
                    financial_persons.update(model.data[name_col].dropna().unique().tolist())
        return sorted(list(financial_persons))

    def _generate_fund_volume_analysis(self, doc: Document, person_name: str, data_models: Dict, analyzers: Dict):
        """生成资金体量分析"""
        doc.add_heading("一、资金体量", level=3)
        
        # 1. 总资金体量
        self._generate_total_fund_volume(doc, person_name, data_models)
        
        # 2. 活跃时间和交易对手
        self._generate_active_time_and_opponents(doc, person_name, data_models)
        
        # 3. 存取现和大额资金
        self._generate_cash_and_large_amounts(doc, person_name, data_models)

    def _generate_total_fund_volume(self, doc: Document, person_name: str, data_models: Dict):
        """生成总资金体量分析"""
        p = doc.add_paragraph()
        p.add_run("1.总资金体量：").bold = True
        
        total_income = 0
        total_expense = 0
        balances = {}
        time_ranges = {}
        
        # 银行数据
        if data_models.get('bank') and not data_models['bank'].data.empty:
            bank_data = data_models['bank'].data[data_models['bank'].data[data_models['bank'].name_column] == person_name]
            if not bank_data.empty:
                income = bank_data[bank_data['交易金额'] > 0]['交易金额'].sum()
                expense = abs(bank_data[bank_data['交易金额'] < 0]['交易金额'].sum())
                total_income += income
                total_expense += expense
                
                # 获取时间范围（格式化为年月）
                if '交易日期' in bank_data.columns:
                    min_date = pd.to_datetime(bank_data['交易日期'].min(), errors='coerce')
                    max_date = pd.to_datetime(bank_data['交易日期'].max(), errors='coerce')
                    if pd.notna(min_date) and pd.notna(max_date):
                        time_ranges['银行'] = f"{min_date.strftime('%Y年%m月')}至{max_date.strftime('%Y年%m月')}"
                
                # 余额统计
                if '账户余额' in bank_data.columns and not bank_data['账户余额'].empty:
                    balances['银行'] = bank_data['账户余额'].iloc[-1]  # 最后一条记录的余额
                
                # 银行总进账和总出账（加粗条件：大于1000万，标题加下划线）
                # 与旧版保持一致，只计算转账交易
                transfer_data = bank_data[bank_data['存取现标识'] == '转账']
                if not transfer_data.empty:
                    income = transfer_data[transfer_data['交易金额'] > 0]['交易金额'].sum()
                    expense = abs(transfer_data[transfer_data['交易金额'] < 0]['交易金额'].sum())
                else:
                    income = 0
                    expense = 0
                
                p.add_run("银行总进账").underline = True
                if income >= 10000000:  # 1000万
                    p.add_run(f"{income:,.2f}元").bold = True
                else:
                    p.add_run(f"{income:,.2f}元")
                p.add_run("、")
                p.add_run("银行总出账").underline = True
                if abs(expense) >= 10000000:  # 1000万
                    p.add_run(f"{expense:,.2f}元").bold = True
                else:
                    p.add_run(f"{expense:,.2f}元")
                p.add_run("；")
                
                # 各银行余额
                if '银行类型' in bank_data.columns:
                    bank_balances = bank_data.groupby('银行类型')['账户余额'].last()
                    for bank_name, balance in bank_balances.items():
                        p.add_run(f"{bank_name}余额{balance:,.2f}元；")
        
        # 微信数据
        if data_models.get('wechat') and not data_models['wechat'].data.empty:
            wechat_data = data_models['wechat'].data[data_models['wechat'].data[data_models['wechat'].name_column] == person_name]
            if not wechat_data.empty:
                income = wechat_data[wechat_data['交易金额'] > 0]['交易金额'].sum()
                expense = abs(wechat_data[wechat_data['交易金额'] < 0]['交易金额'].sum())
                total_income += income
                total_expense += expense
                
                # 获取时间范围（格式化为年月）
                if '交易日期' in wechat_data.columns:
                    min_date = pd.to_datetime(wechat_data['交易日期'].min(), errors='coerce')
                    max_date = pd.to_datetime(wechat_data['交易日期'].max(), errors='coerce')
                    if pd.notna(min_date) and pd.notna(max_date):
                        time_ranges['微信'] = f"{min_date.strftime('%Y年%m月')}至{max_date.strftime('%Y年%m月')}"
                
                # 余额统计
                if '账户余额' in wechat_data.columns and not wechat_data['账户余额'].empty:
                    balances['微信'] = wechat_data['账户余额'].iloc[-1]
                
                # 微信总进账和总出账（标题加下划线）
                p.add_run("微信总进账").underline = True
                if income >= 10000000:  # 1000万
                    p.add_run(f"{income:,.2f}元").bold = True
                else:
                    p.add_run(f"{income:,.2f}元")
                p.add_run("、")
                p.add_run("微信总出账").underline = True
                if abs(expense) >= 10000000:  # 1000万
                    p.add_run(f"{expense:,.2f}元").bold = True
                else:
                    p.add_run(f"{expense:,.2f}元")
                p.add_run("；")
        
        # 支付宝数据
        if data_models.get('alipay') and not data_models['alipay'].data.empty:
            alipay_data = data_models['alipay'].data[data_models['alipay'].data[data_models['alipay'].name_column] == person_name]
            if not alipay_data.empty:
                # 支付宝数据可能没有直接的借贷标识，需要根据交易类型判断
                income = alipay_data[alipay_data['交易金额'] > 0]['交易金额'].sum()
                expense = abs(alipay_data[alipay_data['交易金额'] < 0]['交易金额'].sum())
                total_income += income
                total_expense += expense
                
                # 获取时间范围（格式化为年月）
                if '交易日期' in alipay_data.columns:
                    min_date = pd.to_datetime(alipay_data['交易日期'].min(), errors='coerce')
                    max_date = pd.to_datetime(alipay_data['交易日期'].max(), errors='coerce')
                    if pd.notna(min_date) and pd.notna(max_date):
                        time_ranges['支付宝'] = f"{min_date.strftime('%Y年%m月')}至{max_date.strftime('%Y年%m月')}"
                
                # 支付宝总进账和总出账（标题加下划线）
                p.add_run("支付宝总进账").underline = True
                if income >= 10000000:  # 1000万
                    p.add_run(f"{income:,.2f}元").bold = True
                else:
                    p.add_run(f"{income:,.2f}元")
                p.add_run("、")
                p.add_run("支付宝总出账").underline = True
                if abs(expense) >= 10000000:  # 1000万
                    p.add_run(f"{expense:,.2f}元").bold = True
                else:
                    p.add_run(f"{expense:,.2f}元")
                p.add_run("；")
        
        # 显示时间范围
        time_info = []
        for platform, time_range in time_ranges.items():
            time_info.append(f"{platform}({time_range})")
        if time_info:
            p.add_run("时间跨度：").underline = True
            p.add_run(f"{', '.join(time_info)}；")
        
        # 显示余额（微信和支付宝余额）
        balance_info = []
        if '微信' in balances:
            balance_info.append(f"微信余额{balances['微信']:,.2f}元")
        if '支付宝' in balances:
            balance_info.append(f"支付宝余额{balances['支付宝']:,.2f}元")
        if balance_info:
            p.add_run("余额：").underline = True
            p.add_run(f"{', '.join(balance_info)}；")
        
        # 最常用的银行（如果有银行数据）
        if data_models.get('bank') and not data_models['bank'].data.empty:
            bank_data = data_models['bank'].data[data_models['bank'].data[data_models['bank'].name_column] == person_name]
            if not bank_data.empty and '银行类型' in bank_data.columns:
                bank_counts = bank_data['银行类型'].value_counts().head(3)
                if not bank_counts.empty:
                    bank_info = []
                    for bank_name, count in bank_counts.items():
                        bank_info.append(f"{bank_name}({count}笔)")
                    p.add_run(f"最常用的银行：{', '.join(bank_info)}；")

    def _generate_active_time_and_opponents(self, doc: Document, person_name: str, data_models: Dict):
        """生成活跃时间和交易对手分析"""
        p = doc.add_paragraph()
        p.add_run("2.活跃时间和交易对手：").bold = True
        
        # 各平台活跃年份和交易对手分析
        platforms = ['银行', '微信', '支付宝']
        platform_data_map = {
            '银行': data_models.get('bank'),
            '微信': data_models.get('wechat'),
            '支付宝': data_models.get('alipay')
        }
        
        for platform in platforms:
            model = platform_data_map.get(platform)
            if model and not model.data.empty:
                platform_data = model.data[model.data[model.name_column] == person_name]
                if not platform_data.empty:
                    # 年份分析
                    if '交易日期' in platform_data.columns:
                        # 使用.copy()避免SettingWithCopyWarning
                        platform_data_copy = platform_data.copy()
                        platform_data_copy.loc[:, '年份'] = pd.to_datetime(platform_data_copy['交易日期']).dt.year
                        # 分别计算收入和支出
                        platform_data_copy['收入金额'] = platform_data_copy['交易金额'].apply(lambda x: x if x > 0 else 0)
                        platform_data_copy['支出金额'] = platform_data_copy['交易金额'].apply(lambda x: abs(x) if x < 0 else 0)
                        
                        yearly_stats = platform_data_copy.groupby('年份').agg({
                            '收入金额': 'sum',
                            '支出金额': 'sum',
                            '交易金额': 'count'
                        }).round(2)
                        yearly_stats.columns = ['总进账', '总出账', '交易次数']
                        yearly_stats['总体量'] = yearly_stats['总进账'] + yearly_stats['总出账']
                        
                        # 获取前三名年份
                        top_years = yearly_stats.nlargest(3, '总体量')
                        year_info = []
                        for year, row in top_years.iterrows():
                            year_info.append(f"{year}年（总进账{row['总进账']:,.2f}元、总出账{row['总出账']:,.2f}元，总体量{row['总体量']:,.2f}元）")
                        if year_info:
                            p.add_run(f"{platform}").underline = True
                            p.add_run("资金总量前三的年份：")
                            p.add_run(f"{', '.join(year_info)}；")
                    
                    # 交易对手分析
                    if '对方姓名' in platform_data.columns:
                        # 过滤掉空值和未知的对方姓名
                        opponent_data = platform_data[platform_data['对方姓名'].notna() & 
                                                    (platform_data['对方姓名'] != '') & 
                                                    (platform_data['对方姓名'] != '未知')]
                        if not opponent_data.empty:
                            opponent_stats = opponent_data.groupby('对方姓名')['交易金额'].agg(['sum', 'count']).round(2)
                            opponent_stats['总金额'] = opponent_stats['sum'].abs()
                            opponent_stats = opponent_stats.nlargest(3, '总金额')
                            
                            opponent_info = []
                            for opponent, row in opponent_stats.iterrows():
                                opponent_info.append(f"{opponent}（{row['总金额']:,.2f}元）")
                            if opponent_info:
                                p.add_run(f"{platform}").underline = True
                                p.add_run("交易资金总量的对手前三名：")
                                p.add_run(f"{', '.join(opponent_info)}；")

    def _generate_cash_and_large_amounts(self, doc: Document, person_name: str, data_models: Dict):
        """生成存取现和大额资金分析"""
        p = doc.add_paragraph()
        p.add_run("3.存取现和大额资金：").bold = True
        
        # 存取现分析
        if data_models.get('bank') and not data_models['bank'].data.empty:
            bank_data = data_models['bank'].data[data_models['bank'].data[data_models['bank'].name_column] == person_name]
            if not bank_data.empty:
                # 筛选存现和取现数据
                deposit_data = bank_data[bank_data['存取现标识'] == '存现']
                withdraw_data = bank_data[bank_data['存取现标识'] == '取现']
                
                # 计算存现总额和取现总额
                deposit_total = deposit_data['交易金额'].sum() if not deposit_data.empty else 0
                withdraw_total = abs(withdraw_data['交易金额'].sum()) if not withdraw_data.empty else 0
                
                # 存现总额和取现总额（加粗条件）
                p.add_run("存现总额").underline = True
                deposit_text = f"{deposit_total:,.2f}元"
                withdraw_text = f"{withdraw_total:,.2f}元"
                if deposit_total >= 1000000:  # 100万
                    p.add_run(deposit_text).bold = True
                else:
                    p.add_run(deposit_text)
                p.add_run("，")
                p.add_run("取现总额").underline = True
                if withdraw_total >= 1000000:  # 100万
                    p.add_run(withdraw_text).bold = True
                else:
                    p.add_run(withdraw_text)
                p.add_run("；")
                
                # 单笔存现1万及以上
                large_deposit_data = deposit_data[deposit_data['交易金额'] >= 10000]
                if not large_deposit_data.empty:
                    large_deposit_count = len(large_deposit_data)
                    large_deposit_amount = large_deposit_data['交易金额'].sum()
                    
                    # 按银行统计
                    bank_deposit_stats = large_deposit_data.groupby('银行类型')['交易金额'].agg(['sum', 'count'])
                    bank_deposit_info = []
                    for bank_name, row in bank_deposit_stats.iterrows():
                        # 存现金额大于100万加粗
                        deposit_info = f"{bank_name}{row['sum']:,.2f}元{row['count']}次"
                        if row['sum'] >= 1000000:  # 100万
                            bank_deposit_info.append(deposit_info)
                        else:
                            bank_deposit_info.append(deposit_info)
                    
                    large_deposit_text = f"单笔存现1万及以上的总金额{large_deposit_amount:,.2f}元{large_deposit_count}次（{', '.join(bank_deposit_info)}）"
                    if large_deposit_amount >= 1000000:  # 100万
                        p.add_run(large_deposit_text).bold = True
                    else:
                        p.add_run(large_deposit_text)
                    p.add_run("；")
                
                # 单笔取现1万及以上
                large_withdraw_data = withdraw_data[withdraw_data['交易金额'] <= -10000]
                if not large_withdraw_data.empty:
                    large_withdraw_count = len(large_withdraw_data)
                    large_withdraw_amount = abs(large_withdraw_data['交易金额'].sum())
                    
                    # 按银行统计
                    bank_withdraw_stats = large_withdraw_data.groupby('银行类型')['交易金额'].agg(['sum', 'count'])
                    bank_withdraw_info = []
                    for bank_name, row in bank_withdraw_stats.iterrows():
                        # 取现金额大于100万加粗
                        withdraw_info = f"{bank_name}{abs(row['sum']):,.2f}元{row['count']}次"
                        if abs(row['sum']) >= 1000000:  # 100万
                            bank_withdraw_info.append(withdraw_info)
                        else:
                            bank_withdraw_info.append(withdraw_info)
                    
                    large_withdraw_text = f"单笔取现1万及以上的总金额{large_withdraw_amount:,.2f}元{large_withdraw_count}次（{', '.join(bank_withdraw_info)}）"
                    if large_withdraw_amount >= 1000000:  # 100万
                        p.add_run(large_withdraw_text).bold = True
                    else:
                        p.add_run(large_withdraw_text)
                    p.add_run("；")
                
                # 银行总转账金额（加粗条件：大于1000万）
                transfer_data = bank_data[bank_data['存取现标识'] == '转账']
                if not transfer_data.empty:
                    transfer_amount = transfer_data['交易金额'].abs().sum()
                    
                    # 显示总转账金额
                    p.add_run("银行总转账金额").underline = True
                    transfer_text = f"{transfer_amount:,.2f}元"
                    if transfer_amount >= 10000000:  # 1000万
                        p.add_run(transfer_text).bold = True
                    else:
                        p.add_run(transfer_text)
                    p.add_run("；")
                
                # 单笔转账5万元以上的总金额
                large_transfer_data = transfer_data[transfer_data['交易金额'].abs() >= 50000]
                if not large_transfer_data.empty:
                    large_transfer_count = len(large_transfer_data)
                    large_transfer_amount = large_transfer_data['交易金额'].abs().sum()
                    
                    # 按银行统计
                    bank_transfer_stats = large_transfer_data.groupby('银行类型')['交易金额'].agg(['sum', 'count'])
                    bank_transfer_info = []
                    for bank_name, row in bank_transfer_stats.iterrows():
                        # 转账金额大于1000万加粗
                        transfer_info = f"{bank_name}{abs(row['sum']):,.2f}元{row['count']}次"
                        bank_transfer_info.append(transfer_info)
                    
                    large_transfer_text = f"单笔转账5万元以上的总金额{large_transfer_amount:,.2f}元{large_transfer_count}次（{', '.join(bank_transfer_info)}）"
                    if large_transfer_amount >= 10000000:  # 1000万
                        p.add_run(large_transfer_text).bold = True
                    else:
                        p.add_run(large_transfer_text)
                    p.add_run("；")

    def _generate_special_fund_analysis(self, doc: Document, person_name: str, data_models: Dict, analyzers: Dict):
        """生成特殊资金分析"""
        doc.add_heading("二、特殊资金分析", level=3)
        
        # 1. 纯进、出账统计
        self._generate_pure_income_expense_stats(doc, person_name, data_models)
        
        # 2. 特殊金额统计
        self._generate_special_amount_stats(doc, person_name, data_models, analyzers)
        
        # 3. 特殊日期统计
        self._generate_special_date_stats(doc, person_name, data_models, analyzers)

    def _generate_pure_income_expense_stats(self, doc: Document, person_name: str, data_models: Dict):
        """生成纯进、出账统计"""
        p = doc.add_paragraph()
        p.add_run("1.纯进、出账统计：").bold = True
        
        pure_income_opponents = {}
        pure_expense_opponents = {}
        
        # 分析各平台数据
        platforms = ['银行', '微信', '支付宝']
        platform_data_map = {
            '银行': data_models.get('bank'),
            '微信': data_models.get('wechat'),
            '支付宝': data_models.get('alipay')
        }
        
        for platform in platforms:
            model = platform_data_map.get(platform)
            if model and not model.data.empty:
                platform_data = model.data[model.data[model.name_column] == person_name]
                if not platform_data.empty and '对方姓名' in platform_data.columns:
                    # 过滤掉空值和未知的对方姓名
                    opponent_data = platform_data[platform_data['对方姓名'].notna() & 
                                                (platform_data['对方姓名'] != '') & 
                                                (platform_data['对方姓名'] != '未知')]
                    
                    if not opponent_data.empty:
                        # 计算每个对手方的净流入/流出
                        opponent_net = opponent_data.groupby('对方姓名')['交易金额'].sum()
                        
                        # 纯收入对手（净流入>0）
                        pure_income = opponent_net[opponent_net > 0]
                        for opponent, amount in pure_income.items():
                            if opponent not in pure_income_opponents:
                                pure_income_opponents[opponent] = 0
                            pure_income_opponents[opponent] += amount
                        
                        # 纯支出对手（净流出<0）
                        pure_expense = opponent_net[opponent_net < 0]
                        for opponent, amount in pure_expense.items():
                            if opponent not in pure_expense_opponents:
                                pure_expense_opponents[opponent] = 0
                            pure_expense_opponents[opponent] += abs(amount)
        
        # 统计纯收入和纯支出对手数量
        pure_income_count = len(pure_income_opponents)
        pure_expense_count = len(pure_expense_opponents)
        
        p.add_run("纯收入对手共计").underline = True
        p.add_run(f"{pure_income_count}个，")
        p.add_run("纯支出对手共计").underline = True
        p.add_run(f"{pure_expense_count}个；")
        
        # 列出纯进账前五名
        if pure_income_opponents:
            sorted_income = sorted(pure_income_opponents.items(), key=lambda x: x[1], reverse=True)[:5]
            income_info = [f"{opponent}（{amount:,.2f}元）" for opponent, amount in sorted_income]
            p.add_run(f"纯进账的对手前五名：{', '.join(income_info)}；")
        
        # 列出纯出账前五名
        if pure_expense_opponents:
            sorted_expense = sorted(pure_expense_opponents.items(), key=lambda x: x[1], reverse=True)[:5]
            expense_info = [f"{opponent}（{amount:,.2f}元）" for opponent, amount in sorted_expense]
            p.add_run(f"纯出账的对手前五名：{', '.join(expense_info)}；")

    def _generate_special_amount_stats(self, doc: Document, person_name: str, data_models: Dict, analyzers: Dict):
        """生成特殊金额统计（使用原Word报告的逻辑）"""
        p = doc.add_paragraph()
        p.add_run("2.特殊金额统计：").bold = True
        
        # 收集各平台的特殊金额交易
        special_amount_transactions = []
        
        # 在各平台数据中查找特殊金额交易
        platforms = ['银行', '微信', '支付宝']
        platform_data_map = {
            '银行': data_models.get('bank'),
            '微信': data_models.get('wechat'),
            '支付宝': data_models.get('alipay')
        }
        
        for platform in platforms:
            model = platform_data_map.get(platform)
            if model and not model.data.empty:
                platform_data = model.data[model.data[model.name_column] == person_name]
                if not platform_data.empty:
                    # 使用原Word报告的逻辑来识别特殊金额
                    if platform == '银行':
                        analyzer = analyzers.get('bank')
                        if analyzer:
                            special_amounts_df = analyzer.analyze_special_amounts(platform_data.copy())
                            if not special_amounts_df.empty:
                                # 添加到特殊金额交易列表
                                for _, row in special_amounts_df.iterrows():
                                    special_amount_transactions.append({
                                        '平台': platform,
                                        '金额': abs(row[model.amount_column]),
                                        '对方': row.get(model.opposite_name_column, '未知'),
                                        '日期': row.get(model.date_column, '未知日期')
                                    })
                    else:
                        # 对于微信和支付宝，使用配置中的特殊金额定义
                        special_amounts = self.config.get('analysis.special_amount.amounts', [])
                        if special_amounts and model.amount_column in platform_data.columns:
                            for _, row in platform_data.iterrows():
                                amount = abs(row[model.amount_column])
                                if amount in special_amounts:
                                    special_amount_transactions.append({
                                        '平台': platform,
                                        '金额': amount,
                                        '对方': row.get(model.opposite_name_column, '未知'),
                                        '日期': row.get(model.date_column, '未知日期')
                                    })
        
        if special_amount_transactions:
            # 按金额分组统计
            amount_groups = {}
            for tx in special_amount_transactions:
                amount = tx['金额']
                if amount not in amount_groups:
                    amount_groups[amount] = {'次数': 0, '对手': {}}
                amount_groups[amount]['次数'] += 1
                opponent = tx['对方']
                if opponent not in amount_groups[amount]['对手']:
                    amount_groups[amount]['对手'][opponent] = 0
                amount_groups[amount]['对手'][opponent] += 1
            
            # 使用原Word报告的显示逻辑
            p.add_run(f"{person_name}发生")
            
            # 显示所有特殊金额（不使用"等"字概括）
            unique_amounts = sorted(amount_groups.keys(), reverse=True)
            love_amounts = [520, 1314, 521]
            other_amounts = [amt for amt in unique_amounts if amt not in love_amounts]
            display_amounts = love_amounts + other_amounts
            display_amounts_str = "、".join([f"{amt:.2f}" for amt in display_amounts])
            p.add_run(f"{display_amounts_str}等特殊金额{len(special_amount_transactions)}次，")
            
            # 收集特殊金额信息，按金额分组
            amount_details = {}
            
            # 处理爱情数字（520, 1314, 521）
            love_amounts = [520, 1314, 521]
            for love_amount in love_amounts:
                if love_amount in amount_groups:
                    opponents_info = []
                    for opponent, count in amount_groups[love_amount]['对手'].items():
                        opponents_info.append(f"{opponent}，{count}次")
                    opponents_str = "、".join(opponents_info) if opponents_info else "未知，1次"
                    amount_details[love_amount] = opponents_str
            
            # 处理其他特殊金额
            other_amounts = [amt for amt in unique_amounts if amt not in love_amounts]
            for amount in other_amounts:
                opponents_info = []
                for opponent, count in amount_groups[amount]['对手'].items():
                    opponents_info.append(f"{opponent}，{count}次")
                opponents_str = "、".join(opponents_info) if opponents_info else "未知，1次"
                amount_details[amount] = opponents_str
            
            # 构建显示字符串
            p.add_run("特殊金额包括：")
            
            # 先显示爱情数字
            love_items = []
            for amount in love_amounts:
                if amount in amount_details:
                    love_run = p.add_run(f"{amount:.2f}元（{amount_details[amount]}）")
                    love_run.bold = True
                    love_items.append(amount)
            
            # 再显示其他特殊金额（只显示前三个）
            other_items = [amt for amt in amount_details.keys() if amt not in love_items][:3]
            
            # 如果既有爱情数字又有其他数字，添加分隔符
            if love_items and other_items:
                p.add_run("，以及")
            
            # 添加其他特殊金额
            for i, amount in enumerate(other_items):
                p.add_run(f"{amount:.2f}元（{amount_details[amount]}）")
                if i < len(other_items) - 1:
                    p.add_run("、")
            
            p.add_run("。")
        else:
            p.add_run("未发现特殊金额交易。")

    def _generate_special_date_stats(self, doc: Document, person_name: str, data_models: Dict, analyzers: Dict):
        """生成特殊日期统计（使用原Word报告的逻辑）"""
        p = doc.add_paragraph()
        p.add_run("3.特殊日期统计：").bold = True
        
        # 收集各平台的特殊日期交易
        special_date_transactions = []
        
        # 在各平台数据中查找特殊日期交易
        platforms = ['银行', '微信', '支付宝']
        platform_data_map = {
            '银行': data_models.get('bank'),
            '微信': data_models.get('wechat'),
            '支付宝': data_models.get('alipay')
        }
        
        for platform in platforms:
            model = platform_data_map.get(platform)
            if model and not model.data.empty:
                platform_data = model.data[model.data[model.name_column] == person_name]
                if not platform_data.empty:
                    # 使用原Word报告的逻辑来识别特殊日期
                    if platform == '银行':
                        analyzer = analyzers.get('bank')
                        if analyzer:
                            special_dates_df = analyzer.analyze_special_dates(platform_data.copy())
                            if not special_dates_df.empty:
                                # 添加到特殊日期交易列表
                                for _, row in special_dates_df.iterrows():
                                    special_date_transactions.append({
                                        '平台': platform,
                                        '日期名称': row.get('特殊日期名称', '未知日期'),
                                        '金额': abs(row[model.amount_column]),
                                        '对方': row.get(model.opposite_name_column, '未知')
                                    })
                    else:
                        # 对于微信和支付宝，使用配置中的特殊日期定义
                        special_dates = self.config.get('analysis.special_date.dates', {})
                        if special_dates and '交易日期' in platform_data.columns:
                            # 使用.copy()避免SettingWithCopyWarning
                            platform_data_copy = platform_data.copy()
                            platform_data_copy.loc[:, '交易日期'] = pd.to_datetime(platform_data_copy['交易日期'], errors='coerce')
                            
                            # 查找特殊日期交易
                            for _, row in platform_data_copy.iterrows():
                                tx_date = row['交易日期']
                                if pd.isna(tx_date):
                                    continue
                                    
                                # 检查是否为特殊日期
                                for date_name, date_config in special_dates.items():
                                    if date_config['type'] == 'solar':  # 阳历
                                        if tx_date.month == date_config['month'] and tx_date.day == date_config['day']:
                                            special_date_transactions.append({
                                                '平台': platform,
                                                '日期名称': date_name,
                                                '金额': row.get('交易金额', 0),
                                                '对方': row.get('对方姓名', '未知')
                                            })
        
        if special_date_transactions:
            # 按日期名称分组统计
            date_stats = {}
            for tx in special_date_transactions:
                date_name = tx['日期名称']
                if date_name not in date_stats:
                    date_stats[date_name] = {'次数': 0, '总金额': 0, '对手': set()}
                date_stats[date_name]['次数'] += 1
                date_stats[date_name]['总金额'] += abs(tx['金额'])
                date_stats[date_name]['对手'].add(tx['对方'])
            
            # 按总额排序，取前三名（与旧版Word报告一致）
            date_stats_df = pd.DataFrame([
                {'特殊日期名称': date_name, '总额': stats['总金额'], '次数': stats['次数']}
                for date_name, stats in date_stats.items()
            ])
            date_stats_df['总额'] = date_stats_df['总额'].abs()  # 取绝对值
            top_dates = date_stats_df.nlargest(3, '总额')
            
            date_descriptions = []
            for _, row in top_dates.iterrows():
                date_name = row['特殊日期名称']
                total_amount = row['总额']
                count = row['次数']
                date_descriptions.append(f"{date_name}（总额{total_amount:,.2f}元、{count}次）")
            
            if date_descriptions:
                p.add_run("、".join(date_descriptions) + "。")
            else:
                p.add_run("未发现特殊日期交易。")
        else:
            p.add_run("未发现特殊日期交易。")

    def _generate_key_transactions_analysis(self, doc: Document, person_name: str, data_models: Dict, analyzers: Dict):
        """生成重点收支分析"""
        doc.add_heading("三、重点收支", level=3)
        
        # 初始化重点收支识别引擎
        key_engine = KeyTransactionEngine(self.config)
        
        # 1. 工作收支
        self._generate_work_income_expense(doc, person_name, data_models, key_engine)
        
        # 2. 房产车辆收支
        self._generate_property_vehicle_income_expense(doc, person_name, data_models, key_engine)
        
        # 3. 理财收入
        self._generate_financial_income(doc, person_name, data_models, key_engine)

    def _generate_work_income_expense(self, doc: Document, person_name: str, data_models: Dict, key_engine: KeyTransactionEngine):
        """生成工作收支分析"""
        work_stats = {
            'income_total': 0,
            'income_count': 0,
            'expense_total': 0,
            'expense_count': 0,
            'work_units': set(),
            'time_range': None
        }
        
        # 收集各平台的工作相关交易
        platforms = ['银行', '微信', '支付宝']
        platform_data_map = {
            '银行': data_models.get('bank'),
            '微信': data_models.get('wechat'),
            '支付宝': data_models.get('alipay')
        }
        
        all_work_transactions = []
        all_dates = []
        
        for platform in platforms:
            model = platform_data_map.get(platform)
            if model and not model.data.empty:
                platform_data = model.data[model.data[model.name_column] == person_name]
                if not platform_data.empty:
                    # 识别重点交易
                    if platform == '银行':
                        identified_data = key_engine.identify_key_transactions(
                            platform_data,
                            model.summary_column if hasattr(model, 'summary_column') else None,
                            model.remark_column if hasattr(model, 'remark_column') else None,
                            model.type_column if hasattr(model, 'type_column') else None,
                            model.amount_column,
                            model.opposite_name_column if hasattr(model, 'opposite_name_column') else None
                        )
                    else:
                        # 微信和支付宝使用不同的列名
                        identified_data = key_engine.identify_key_transactions(
                            platform_data,
                            None,  # 微信支付宝没有摘要列
                            model.remark_column if hasattr(model, 'remark_column') else None,
                            model.type_column if hasattr(model, 'type_column') else None,
                            model.amount_column,
                            model.opposite_name_column if hasattr(model, 'opposite_name_column') else None
                        )
                    
                    # 筛选工作收入相关交易
                    work_income_data = identified_data[identified_data['是否工作收入']]
                    if not work_income_data.empty:
                        work_stats['income_total'] += work_income_data[model.amount_column].sum()
                        work_stats['income_count'] += len(work_income_data)
                        
                        # 收集工作单位信息
                        if '对方姓名' in work_income_data.columns:
                            work_stats['work_units'].update(
                                work_income_data['对方姓名'].dropna().unique()
                            )
                        
                        # 收集日期信息
                        if '交易日期' in work_income_data.columns:
                            all_dates.extend(work_income_data['交易日期'].tolist())
                        
                        all_work_transactions.append(work_income_data)
        
        # 如果有工作收入数据，则显示
        if work_stats['income_total'] > 0 or work_stats['expense_total'] > 0:
            p = doc.add_paragraph()
            p.add_run("1.工作收支：").bold = True
            
            if work_stats['income_total'] > 0:
                # 计算时间范围
                if all_dates:
                    dates = pd.to_datetime([d for d in all_dates if not pd.isna(d)])
                    if len(dates) > 0:
                        time_range = f"{dates.min().strftime('%Y年%m月')}至{dates.max().strftime('%Y年%m月')}"
                        work_stats['time_range'] = time_range
                
                income_info = f"工作收入{work_stats['income_total']:,.2f}元"
                if work_stats['time_range']:
                    income_info += f"（{work_stats['time_range']}，{work_stats['income_count']}次）"
                
                p.add_run("工作收入").underline = True
                p.add_run(f"{work_stats['income_total']:,.2f}元")
                if work_stats['time_range']:
                    p.add_run(f"（{work_stats['time_range']}，{work_stats['income_count']}次）")
                
                # 工作单位信息
                if work_stats['work_units']:
                    p.add_run("，疑似工作单位有")
                    p.add_run(f"{len(work_stats['work_units'])}个").underline = True
                    units = list(work_stats['work_units'])[:3]  # 最多显示3个
                    p.add_run(f"（{', '.join(units)}）")
                p.add_run("；")

    def _generate_property_vehicle_income_expense(self, doc: Document, person_name: str, data_models: Dict, key_engine: KeyTransactionEngine):
        """生成房产车辆收支分析"""
        property_stats = {
            'income_total': 0,
            'income_count': 0,
            'expense_total': 0,
            'expense_count': 0,
            'time_range': None,
            'income_opponents': {},  # 新增：记录收入对手信息
            'expense_opponents': {}  # 新增：记录支出对手信息
        }
        
        vehicle_stats = {
            'income_total': 0,
            'income_count': 0,
            'expense_total': 0,
            'expense_count': 0,
            'time_range': None,
            'income_opponents': {},  # 新增：记录收入对手信息
            'expense_opponents': {}  # 新增：记录支出对手信息
        }
        
        rental_stats = {
            'income_total': 0,
            'income_count': 0,
            'expense_total': 0,
            'expense_count': 0,
            'time_range': None,
            'income_opponents': {},  # 新增：记录收入对手信息
            'expense_opponents': {}  # 新增：记录支出对手信息
        }
        
        # 收集各平台的房产车辆相关交易
        platforms = ['银行', '微信', '支付宝']
        platform_data_map = {
            '银行': data_models.get('bank'),
            '微信': data_models.get('wechat'),
            '支付宝': data_models.get('alipay')
        }
        
        all_property_dates = []
        all_vehicle_dates = []
        all_rental_dates = []
        
        for platform in platforms:
            model = platform_data_map.get(platform)
            if model and not model.data.empty:
                platform_data = model.data[model.data[model.name_column] == person_name]
                if not platform_data.empty:
                    # 识别重点交易
                    if platform == '银行':
                        identified_data = key_engine.identify_key_transactions(
                            platform_data,
                            model.summary_column if hasattr(model, 'summary_column') else None,
                            model.remark_column if hasattr(model, 'remark_column') else None,
                            model.type_column if hasattr(model, 'type_column') else None,
                            model.amount_column,
                            model.opposite_name_column if hasattr(model, 'opposite_name_column') else None
                        )
                    else:
                        # 微信和支付宝使用不同的列名
                        identified_data = key_engine.identify_key_transactions(
                            platform_data,
                            None,  # 微信支付宝没有摘要列
                            model.remark_column if hasattr(model, 'remark_column') else None,
                            model.type_column if hasattr(model, 'type_column') else None,
                            model.amount_column,
                            model.opposite_name_column if hasattr(model, 'opposite_name_column') else None
                        )
                    
                    # 房产相关交易
                    property_income_data = identified_data[identified_data['是否房产收入']]
                    property_expense_data = identified_data[identified_data['是否房产支出']]
                    
                    if not property_income_data.empty:
                        property_stats['income_total'] += property_income_data[model.amount_column].sum()
                        property_stats['income_count'] += len(property_income_data)
                        if '交易日期' in property_income_data.columns:
                            all_property_dates.extend(property_income_data['交易日期'].tolist())
                        # 记录对手信息
                        if '对方姓名' in property_income_data.columns:
                            for _, row in property_income_data.iterrows():
                                opponent = row['对方姓名']
                                amount = row[model.amount_column]
                                if opponent not in property_stats['income_opponents']:
                                    property_stats['income_opponents'][opponent] = 0
                                property_stats['income_opponents'][opponent] += amount
                    
                    if not property_expense_data.empty:
                        property_stats['expense_total'] += abs(property_expense_data[model.amount_column].sum())
                        property_stats['expense_count'] += len(property_expense_data)
                        if '交易日期' in property_expense_data.columns:
                            all_property_dates.extend(property_expense_data['交易日期'].tolist())
                        # 记录对手信息
                        if '对方姓名' in property_expense_data.columns:
                            for _, row in property_expense_data.iterrows():
                                opponent = row['对方姓名']
                                amount = abs(row[model.amount_column])
                                if opponent not in property_stats['expense_opponents']:
                                    property_stats['expense_opponents'][opponent] = 0
                                property_stats['expense_opponents'][opponent] += amount
                    
                    # 车辆相关交易
                    vehicle_income_data = identified_data[identified_data['是否车辆收入']]
                    vehicle_expense_data = identified_data[identified_data['是否车辆支出']]
                    
                    if not vehicle_income_data.empty:
                        vehicle_stats['income_total'] += vehicle_income_data[model.amount_column].sum()
                        vehicle_stats['income_count'] += len(vehicle_income_data)
                        if '交易日期' in vehicle_income_data.columns:
                            all_vehicle_dates.extend(vehicle_income_data['交易日期'].tolist())
                        # 记录对手信息
                        if '对方姓名' in vehicle_income_data.columns:
                            for _, row in vehicle_income_data.iterrows():
                                opponent = row['对方姓名']
                                amount = row[model.amount_column]
                                if opponent not in vehicle_stats['income_opponents']:
                                    vehicle_stats['income_opponents'][opponent] = 0
                                vehicle_stats['income_opponents'][opponent] += amount
                    
                    if not vehicle_expense_data.empty:
                        vehicle_stats['expense_total'] += abs(vehicle_expense_data[model.amount_column].sum())
                        vehicle_stats['expense_count'] += len(vehicle_expense_data)
                        if '交易日期' in vehicle_expense_data.columns:
                            all_vehicle_dates.extend(vehicle_expense_data['交易日期'].tolist())
                        # 记录对手信息
                        if '对方姓名' in vehicle_expense_data.columns:
                            for _, row in vehicle_expense_data.iterrows():
                                opponent = row['对方姓名']
                                amount = abs(row[model.amount_column])
                                if opponent not in vehicle_stats['expense_opponents']:
                                    vehicle_stats['expense_opponents'][opponent] = 0
                                vehicle_stats['expense_opponents'][opponent] += amount
                    
                    # 租金相关交易
                    rental_income_data = identified_data[identified_data['是否租金收入']]
                    rental_expense_data = identified_data[identified_data['是否租金支出']]
                    
                    if not rental_income_data.empty:
                        rental_stats['income_total'] += rental_income_data[model.amount_column].sum()
                        rental_stats['income_count'] += len(rental_income_data)
                        if '交易日期' in rental_income_data.columns:
                            all_rental_dates.extend(rental_income_data['交易日期'].tolist())
                        # 记录对手信息
                        if '对方姓名' in rental_income_data.columns:
                            for _, row in rental_income_data.iterrows():
                                opponent = row['对方姓名']
                                amount = row[model.amount_column]
                                if opponent not in rental_stats['income_opponents']:
                                    rental_stats['income_opponents'][opponent] = 0
                                rental_stats['income_opponents'][opponent] += amount
                    
                    if not rental_expense_data.empty:
                        rental_stats['expense_total'] += abs(rental_expense_data[model.amount_column].sum())
                        rental_stats['expense_count'] += len(rental_expense_data)
                        if '交易日期' in rental_expense_data.columns:
                            all_rental_dates.extend(rental_expense_data['交易日期'].tolist())
                        # 记录对手信息
                        if '对方姓名' in rental_expense_data.columns:
                            for _, row in rental_expense_data.iterrows():
                                opponent = row['对方姓名']
                                amount = abs(row[model.amount_column])
                                if opponent not in rental_stats['expense_opponents']:
                                    rental_stats['expense_opponents'][opponent] = 0
                                rental_stats['expense_opponents'][opponent] += amount
        
        # 如果有任何房产车辆相关数据，则显示
        if (property_stats['income_total'] > 0 or property_stats['expense_total'] > 0 or
            vehicle_stats['income_total'] > 0 or vehicle_stats['expense_total'] > 0 or
            rental_stats['income_total'] > 0 or rental_stats['expense_total'] > 0):
            
            p = doc.add_paragraph()
            p.add_run("2.房产车辆收支：").bold = True
            
            # 房产相关
            if property_stats['income_total'] > 0 or property_stats['expense_total'] > 0:
                property_info = f"涉及房产收入{property_stats['income_total']:,.2f}元{property_stats['income_count']}次"
                if property_stats['expense_total'] > 0:
                    property_info += f"，房产支出{property_stats['expense_total']:,.2f}元{property_stats['expense_count']}次"
                
                # 时间范围
                if all_property_dates:
                    dates = pd.to_datetime([d for d in all_property_dates if not pd.isna(d)])
                    if len(dates) > 0:
                        time_range = f"{dates.min().strftime('%Y年')}、{dates.max().strftime('%Y年')}"
                        property_info += f"，时间是{time_range}"
                
                # 对手信息
                opponent_info = []
                if property_stats['income_opponents']:
                    income_opponents = sorted(property_stats['income_opponents'].items(), key=lambda x: x[1], reverse=True)[:3]
                    opponent_info.append("收入对手：" + "、".join([f"{opponent}（{amount:,.0f}元）" for opponent, amount in income_opponents]))
                if property_stats['expense_opponents']:
                    expense_opponents = sorted(property_stats['expense_opponents'].items(), key=lambda x: x[1], reverse=True)[:3]
                    opponent_info.append("支出对手：" + "、".join([f"{opponent}（{amount:,.0f}元）" for opponent, amount in expense_opponents]))
                
                if opponent_info:
                    property_info += "，交易对手列举" + "；".join(opponent_info)
                
                p.add_run("房产收入").underline = True
                p.add_run(f"{property_stats['income_total']:,.2f}元{property_stats['income_count']}次")
                if property_stats['expense_total'] > 0:
                    p.add_run("，房产支出").underline = True
                    p.add_run(f"{property_stats['expense_total']:,.2f}元{property_stats['expense_count']}次")
                
                # 时间范围
                if all_property_dates:
                    dates = pd.to_datetime([d for d in all_property_dates if not pd.isna(d)])
                    if len(dates) > 0:
                        time_range = f"{dates.min().strftime('%Y年')}、{dates.max().strftime('%Y年')}"
                        p.add_run(f"，时间是{time_range}")
                
                # 对手信息
                opponent_info = []
                if property_stats['income_opponents']:
                    income_opponents = sorted(property_stats['income_opponents'].items(), key=lambda x: x[1], reverse=True)[:3]
                    opponent_info.append("收入对手：" + "、".join([f"{opponent}（{amount:,.0f}元）" for opponent, amount in income_opponents]))
                if property_stats['expense_opponents']:
                    expense_opponents = sorted(property_stats['expense_opponents'].items(), key=lambda x: x[1], reverse=True)[:3]
                    opponent_info.append("支出对手：" + "、".join([f"{opponent}（{amount:,.0f}元）" for opponent, amount in expense_opponents]))
                
                if opponent_info:
                    p.add_run("，交易对手列举" + "；".join(opponent_info))
                p.add_run("；")
            
            # 车辆相关
            if vehicle_stats['income_total'] > 0 or vehicle_stats['expense_total'] > 0:
                vehicle_info = f"涉及车辆收入{vehicle_stats['income_total']:,.2f}元{vehicle_stats['income_count']}次"
                if vehicle_stats['expense_total'] > 0:
                    vehicle_info += f"，车辆支出{vehicle_stats['expense_total']:,.2f}元{vehicle_stats['expense_count']}次"
                
                # 时间范围
                if all_vehicle_dates:
                    dates = pd.to_datetime([d for d in all_vehicle_dates if not pd.isna(d)])
                    if len(dates) > 0:
                        time_range = f"{dates.min().strftime('%Y年')}、{dates.max().strftime('%Y年')}"
                        vehicle_info += f"，时间是{time_range}"
                
                # 对手信息
                opponent_info = []
                if vehicle_stats['income_opponents']:
                    income_opponents = sorted(vehicle_stats['income_opponents'].items(), key=lambda x: x[1], reverse=True)[:3]
                    opponent_info.append("收入对手：" + "、".join([f"{opponent}（{amount:,.0f}元）" for opponent, amount in income_opponents]))
                if vehicle_stats['expense_opponents']:
                    expense_opponents = sorted(vehicle_stats['expense_opponents'].items(), key=lambda x: x[1], reverse=True)[:3]
                    opponent_info.append("支出对手：" + "、".join([f"{opponent}（{amount:,.0f}元）" for opponent, amount in expense_opponents]))
                
                if opponent_info:
                    vehicle_info += "，交易对手列举" + "；".join(opponent_info)
                
                p.add_run("车辆收入").underline = True
                p.add_run(f"{vehicle_stats['income_total']:,.2f}元{vehicle_stats['income_count']}次")
                if vehicle_stats['expense_total'] > 0:
                    p.add_run("，车辆支出").underline = True
                    p.add_run(f"{vehicle_stats['expense_total']:,.2f}元{vehicle_stats['expense_count']}次")
                
                # 时间范围
                if all_vehicle_dates:
                    dates = pd.to_datetime([d for d in all_vehicle_dates if not pd.isna(d)])
                    if len(dates) > 0:
                        time_range = f"{dates.min().strftime('%Y年')}、{dates.max().strftime('%Y年')}"
                        p.add_run(f"，时间是{time_range}")
                
                # 对手信息
                opponent_info = []
                if vehicle_stats['income_opponents']:
                    income_opponents = sorted(vehicle_stats['income_opponents'].items(), key=lambda x: x[1], reverse=True)[:3]
                    opponent_info.append("收入对手：" + "、".join([f"{opponent}（{amount:,.0f}元）" for opponent, amount in income_opponents]))
                if vehicle_stats['expense_opponents']:
                    expense_opponents = sorted(vehicle_stats['expense_opponents'].items(), key=lambda x: x[1], reverse=True)[:3]
                    opponent_info.append("支出对手：" + "、".join([f"{opponent}（{amount:,.0f}元）" for opponent, amount in expense_opponents]))
                
                if opponent_info:
                    p.add_run("，交易对手列举" + "；".join(opponent_info))
                p.add_run("；")
            
            # 租金相关
            if rental_stats['income_total'] > 0 or rental_stats['expense_total'] > 0:
                rental_info = f"涉及租金收入{rental_stats['income_total']:,.2f}元{rental_stats['income_count']}次"
                if rental_stats['expense_total'] > 0:
                    rental_info += f"，租金支出{rental_stats['expense_total']:,.2f}元{rental_stats['expense_count']}次"
                
                # 时间范围
                if all_rental_dates:
                    dates = pd.to_datetime([d for d in all_rental_dates if not pd.isna(d)])
                    if len(dates) > 0:
                        time_range = f"{dates.min().strftime('%Y年')}、{dates.max().strftime('%Y年')}"
                        rental_info += f"，时间是{time_range}"
                
                # 对手信息
                opponent_info = []
                if rental_stats['income_opponents']:
                    income_opponents = sorted(rental_stats['income_opponents'].items(), key=lambda x: x[1], reverse=True)[:3]
                    opponent_info.append("收入对手：" + "、".join([f"{opponent}（{amount:,.0f}元）" for opponent, amount in income_opponents]))
                if rental_stats['expense_opponents']:
                    expense_opponents = sorted(rental_stats['expense_opponents'].items(), key=lambda x: x[1], reverse=True)[:3]
                    opponent_info.append("支出对手：" + "、".join([f"{opponent}（{amount:,.0f}元）" for opponent, amount in expense_opponents]))
                
                if opponent_info:
                    rental_info += "，交易对手列举" + "；".join(opponent_info)
                
                p.add_run("租金收入").underline = True
                p.add_run(f"{rental_stats['income_total']:,.2f}元{rental_stats['income_count']}次")
                if rental_stats['expense_total'] > 0:
                    p.add_run("，租金支出").underline = True
                    p.add_run(f"{rental_stats['expense_total']:,.2f}元{rental_stats['expense_count']}次")
                
                # 时间范围
                if all_rental_dates:
                    dates = pd.to_datetime([d for d in all_rental_dates if not pd.isna(d)])
                    if len(dates) > 0:
                        time_range = f"{dates.min().strftime('%Y年')}、{dates.max().strftime('%Y年')}"
                        p.add_run(f"，时间是{time_range}")
                
                # 对手信息
                opponent_info = []
                if rental_stats['income_opponents']:
                    income_opponents = sorted(rental_stats['income_opponents'].items(), key=lambda x: x[1], reverse=True)[:3]
                    opponent_info.append("收入对手：" + "、".join([f"{opponent}（{amount:,.0f}元）" for opponent, amount in income_opponents]))
                if rental_stats['expense_opponents']:
                    expense_opponents = sorted(rental_stats['expense_opponents'].items(), key=lambda x: x[1], reverse=True)[:3]
                    opponent_info.append("支出对手：" + "、".join([f"{opponent}（{amount:,.0f}元）" for opponent, amount in expense_opponents]))
                
                if opponent_info:
                    p.add_run("，交易对手列举" + "；".join(opponent_info))
                p.add_run("；")

    def _generate_financial_income(self, doc: Document, person_name: str, data_models: Dict, key_engine: KeyTransactionEngine):
        """生成理财收入分析"""
        financial_stats = {
            'income_total': 0,
            'income_securities': 0,
            'income_securities_count': 0,
            'income_other': 0,
            'income_other_count': 0,
            'expense_total': 0,
            'expense_securities': 0,
            'expense_securities_count': 0,
            'expense_other': 0,
            'expense_other_count': 0
        }
        
        # 收集各平台的理财相关交易
        platforms = ['银行', '微信', '支付宝']
        platform_data_map = {
            '银行': data_models.get('bank'),
            '微信': data_models.get('wechat'),
            '支付宝': data_models.get('alipay')
        }
        
        for platform in platforms:
            model = platform_data_map.get(platform)
            if model and not model.data.empty:
                platform_data = model.data[model.data[model.name_column] == person_name]
                if not platform_data.empty:
                    # 识别重点交易
                    if platform == '银行':
                        identified_data = key_engine.identify_key_transactions(
                            platform_data,
                            model.summary_column if hasattr(model, 'summary_column') else None,
                            model.remark_column if hasattr(model, 'remark_column') else None,
                            model.type_column if hasattr(model, 'type_column') else None,
                            model.amount_column,
                            model.opposite_name_column if hasattr(model, 'opposite_name_column') else None
                        )
                    else:
                        # 微信和支付宝使用不同的列名
                        identified_data = key_engine.identify_key_transactions(
                            platform_data,
                            None,  # 微信支付宝没有摘要列
                            model.remark_column if hasattr(model, 'remark_column') else None,
                            model.type_column if hasattr(model, 'type_column') else None,
                            model.amount_column,
                            model.opposite_name_column if hasattr(model, 'opposite_name_column') else None
                        )
                    
                    # 理财收入相关交易
                    securities_income_data = identified_data[identified_data['是否证券收入']]
                    # 注意：可能没有'是否资产收入'列，需要检查列是否存在
                    if '是否资产收入' in identified_data.columns:
                        other_income_data = identified_data[(identified_data['是否资产收入']) & 
                                                           (~identified_data['是否证券收入']) & 
                                                           (~identified_data['是否房产收入']) & 
                                                           (~identified_data['是否车辆收入']) & 
                                                           (~identified_data['是否租金收入'])]
                    else:
                        # 如果没有'是否资产收入'列，创建一个空的DataFrame
                        other_income_data = pd.DataFrame(columns=identified_data.columns)
                    
                    if not securities_income_data.empty:
                        financial_stats['income_securities'] += securities_income_data[model.amount_column].sum()
                        financial_stats['income_securities_count'] += len(securities_income_data)
                    
                    if not other_income_data.empty:
                        financial_stats['income_other'] += other_income_data[model.amount_column].sum()
                        financial_stats['income_other_count'] += len(other_income_data)
                    
                    financial_stats['income_total'] = financial_stats['income_securities'] + financial_stats['income_other']
                    
                    # 理财支出相关交易
                    securities_expense_data = identified_data[identified_data['是否证券支出']]
                    # 注意：可能没有'是否资产支出'列，需要检查列是否存在
                    if '是否资产支出' in identified_data.columns:
                        other_expense_data = identified_data[(identified_data['是否资产支出']) & 
                                                            (~identified_data['是否证券支出']) & 
                                                            (~identified_data['是否房产支出']) & 
                                                            (~identified_data['是否车辆支出']) & 
                                                            (~identified_data['是否租金支出'])]
                    else:
                        # 如果没有'是否资产支出'列，创建一个空的DataFrame
                        other_expense_data = pd.DataFrame(columns=identified_data.columns)
                    
                    if not securities_expense_data.empty:
                        financial_stats['expense_securities'] += abs(securities_expense_data[model.amount_column].sum())
                        financial_stats['expense_securities_count'] += len(securities_expense_data)
                    
                    if not other_expense_data.empty:
                        financial_stats['expense_other'] += abs(other_expense_data[model.amount_column].sum())
                        financial_stats['expense_other_count'] += len(other_expense_data)
                    
                    financial_stats['expense_total'] = financial_stats['expense_securities'] + financial_stats['expense_other']
        
        # 如果有理财相关数据，则显示
        if financial_stats['income_total'] > 0 or financial_stats['expense_total'] > 0:
            p = doc.add_paragraph()
            p.add_run("3.理财收入：").bold = True
            
            financial_info = []
            
            if financial_stats['income_total'] > 0:
                income_info = f"涉及理财收入{financial_stats['income_total']:,.2f}元"
                income_details = []
                if financial_stats['income_securities'] > 0:
                    income_details.append(f"证券{financial_stats['income_securities']:,.2f}元{financial_stats['income_securities_count']}次")
                if financial_stats['income_other'] > 0:
                    income_details.append(f"其他（非证券）{financial_stats['income_other']:,.2f}元{financial_stats['income_other_count']}次")
                if income_details:
                    income_info += f"，其中{'、'.join(income_details)}"
                financial_info.append(income_info)
            
            if financial_stats['expense_total'] > 0:
                expense_info = f"理财支出{financial_stats['expense_total']:,.2f}元"
                expense_details = []
                if financial_stats['expense_securities'] > 0:
                    expense_details.append(f"证券{financial_stats['expense_securities']:,.2f}元{financial_stats['expense_securities_count']}次")
                if financial_stats['expense_other'] > 0:
                    expense_details.append(f"其他（非证券）{financial_stats['expense_other']:,.2f}元{financial_stats['expense_other_count']}次")
                if expense_details:
                    expense_info += f"，其中{'、'.join(expense_details)}"
                financial_info.append(expense_info)
            
            if financial_stats['income_total'] > 0:
                p.add_run("理财收入").underline = True
                p.add_run(f"{financial_stats['income_total']:,.2f}元")
                income_details = []
                if financial_stats['income_securities'] > 0:
                    income_details.append(f"证券{financial_stats['income_securities']:,.2f}元{financial_stats['income_securities_count']}次")
                if financial_stats['income_other'] > 0:
                    income_details.append(f"其他（非证券）{financial_stats['income_other']:,.2f}元{financial_stats['income_other_count']}次")
                if income_details:
                    p.add_run(f"，其中{'、'.join(income_details)}")
            
            if financial_stats['expense_total'] > 0:
                if financial_stats['income_total'] > 0:
                    p.add_run("，")
                p.add_run("理财支出").underline = True
                p.add_run(f"{financial_stats['expense_total']:,.2f}元")
                expense_details = []
                if financial_stats['expense_securities'] > 0:
                    expense_details.append(f"证券{financial_stats['expense_securities']:,.2f}元{financial_stats['expense_securities_count']}次")
                if financial_stats['expense_other'] > 0:
                    expense_details.append(f"其他（非证券）{financial_stats['expense_other']:,.2f}元{financial_stats['expense_other_count']}次")
                if expense_details:
                    p.add_run(f"，其中{'、'.join(expense_details)}")
            p.add_run("；")

    def _generate_key_persons_analysis(self, doc: Document, data_models: Dict, analyzers: Dict):
        """生成重点人员分析（参考Excel报告逻辑，性能优化版本）"""
        doc.add_heading("四、重点人员", level=3)
        
        # 性能优化：预计算和缓存数据
        print("    预计算重点人员数据...")
        self._precompute_key_persons_data(data_models)
        
        # 1. 存取现与话单匹配的人员
        print("    生成存取现与话单匹配分析...")
        self._generate_cash_call_matching_persons(doc, data_models)
        
        # 2. 大额资金跟踪与话单匹配的人员
        print("    生成大额资金跟踪分析...")
        self._generate_large_fund_call_matching_persons(doc, data_models)
        
        # 3. 大额资金跟踪层级区分的重点人员
        print("    生成重点人员层级分析...")
        self._generate_large_fund_tracking_persons(doc, data_models)

    def _precompute_key_persons_data(self, data_models: Dict):
        """预计算重点人员分析所需的数据，提升性能"""
        # 创建缓存属性
        if not hasattr(self, '_cached_data'):
            self._cached_data = {}
        
        # 预计算存取现数据
        if 'cash_data' not in self._cached_data:
            cash_data = {}
            for data_type in ['bank', 'wechat', 'alipay']:
                if data_models.get(data_type) and not data_models[data_type].data.empty:
                    data = data_models[data_type].data
                    if '存取现标识' in data.columns:
                        cash_records = data[data['存取现标识'].isin(['存现', '取现'])]
                        if not cash_records.empty:
                            cash_data[data_type] = cash_records
            self._cached_data['cash_data'] = cash_data
        
        # 预计算话单数据
        if 'call_data' not in self._cached_data:
            call_data = None
            if data_models.get('call') and not data_models['call'].data.empty:
                call_data = data_models['call'].data.copy()
                # 优化：预转换日期列
                if '呼叫日期' in call_data.columns:
                    call_data['呼叫日期'] = pd.to_datetime(call_data['呼叫日期'], errors='coerce', format='mixed')
                    call_data['date_key'] = call_data['呼叫日期'].dt.date
            self._cached_data['call_data'] = call_data
        
        # 预计算大额交易数据
        if 'large_amount_data' not in self._cached_data:
            large_amount_data = {}
            for data_type in ['bank', 'wechat', 'alipay']:
                if data_models.get(data_type) and not data_models[data_type].data.empty:
                    data = data_models[data_type].data
                    # 筛选大额交易（金额大于5万）
                    large_amounts = data[
                        (data.get('收入金额', 0) > 50000) | 
                        (data.get('支出金额', 0) > 50000)
                    ]
                    if not large_amounts.empty:
                        large_amount_data[data_type] = large_amounts
            self._cached_data['large_amount_data'] = large_amount_data

    def _generate_cash_call_matching_persons(self, doc: Document, data_models: Dict):
        """生成存取现与话单匹配的人员（性能优化版本）"""
        p = doc.add_paragraph()
        p.add_run("1.存取现与话单匹配的人员：").bold = True
        
        try:
            # 使用预计算的数据
            cash_data = self._cached_data.get('cash_data', {})
            call_data = self._cached_data.get('call_data')
            
            if not cash_data or call_data is None:
                p.add_run("无存取现数据或话单数据")
                return
            
            # 性能优化：批量处理存取现与话单匹配
            matching_results = self._batch_cash_call_matching(cash_data, call_data)
            
            if matching_results:
                p.add_run(f"共找到{len(matching_results)}个匹配记录：")
                for i, (person, match_info) in enumerate(matching_results.items()):
                    if i < 10:  # 只显示前10个
                        p.add_run(f"{person}({match_info})")
                        if i < min(9, len(matching_results) - 1):
                            p.add_run("；")
                if len(matching_results) > 10:
                    p.add_run(f"等{len(matching_results)}人")
            else:
                p.add_run("无匹配记录")
                
        except Exception as e:
            p.add_run(f"分析出错: {str(e)}")

    def _batch_cash_call_matching(self, cash_data: Dict, call_data: pd.DataFrame) -> Dict:
        """批量处理存取现与话单匹配，提升性能"""
        matching_results = {}
        
        # 创建话单日期索引
        call_date_groups = call_data.groupby('date_key')
        
        for data_type, cash_records in cash_data.items():
            # 批量处理每个数据源的存取现记录
            for _, cash_record in cash_records.iterrows():
                person_name = cash_record['本方姓名']
                tx_date = pd.to_datetime(cash_record.get('交易日期', ''), errors='coerce')
                
                if pd.isna(tx_date):
                    continue
                
                date_key = tx_date.date()
                if date_key in call_date_groups.groups:
                    same_day_calls = call_date_groups.get_group(date_key)
                    
                    # 查找匹配的话单记录
                    person_calls = same_day_calls[
                        (same_day_calls['本方姓名'] == person_name) | 
                        (same_day_calls['对方姓名'] == person_name)
                    ]
                    
                    if not person_calls.empty:
                        # 收集联系人
                        contacts = set()
                        for _, call in person_calls.iterrows():
                            if call['本方姓名'] == person_name:
                                contacts.add(call['对方姓名'])
                            else:
                                contacts.add(call['本方姓名'])
                        
                        contacts.discard(person_name)  # 移除本人
                        if contacts:
                            if person_name not in matching_results:
                                matching_results[person_name] = []
                            matching_results[person_name].extend(list(contacts))
        
        # 去重并格式化结果
        formatted_results = {}
        for person, contacts in matching_results.items():
            unique_contacts = list(set(contacts))
            formatted_results[person] = f"{len(unique_contacts)}个联系人"
        
        return formatted_results

    def _generate_cash_call_matching_persons_original(self, doc: Document, data_models: Dict):
        """生成存取现与话单匹配的人员（原始版本，保留作为备份）"""
        p = doc.add_paragraph()
        p.add_run("1.存取现与话单匹配的人员：").bold = True
        
        try:
            # 使用Excel报告中的逻辑来获取存取现与话单匹配的信息
            from src.export.excel_exporter import ExcelExporter
            excel_exporter = ExcelExporter()
            
            # 获取存取现与话单匹配的结果
            cash_call_results = excel_exporter._analyze_cash_call_matching(
                data_models, min_amount=10000  # 使用默认的10000元阈值
            )
            
            if not cash_call_results.empty and '话单匹配' in cash_call_results.columns:
                # 筛选出有话单匹配的记录
                matched_records = cash_call_results[cash_call_results['话单匹配'] != '']
                
                if not matched_records.empty:
                    # 收集所有有通联的人名和单位情况
                    matched_persons = set()
                    for _, row in matched_records.iterrows():
                        person_name = row.get('核心人员', '未知')
                        match_info = row.get('话单匹配', '')
                        
                        # 从话单匹配信息中提取对方人员
                        if '对方' in match_info:
                            # 提取对方人员信息
                            parts = match_info.split('对方')
                            if len(parts) > 1:
                                opponents = parts[1].split(',')
                                for opponent in opponents:
                                    opponent = opponent.strip()
                                    if opponent:
                                        matched_persons.add(opponent)
                        
                        # 添加核心人员本身
                        if person_name and person_name != '未知':
                            matched_persons.add(person_name)
                    
                    # 完全列举所有匹配的人员信息
                    if matched_persons:
                        person_list = sorted(list(matched_persons))
                        person_with_units = []
                        for person in person_list:
                            unit_info = self._extract_unit_info_from_call_data(person, data_models)
                            if unit_info:
                                person_with_units.append(f"{person}（{unit_info}）")
                            else:
                                person_with_units.append(person)
                        p.add_run(f"存取现当天有通联的有{', '.join(person_with_units)}；")
                    else:
                        p.add_run("未发现存取现与话单匹配的人员；")
                else:
                    p.add_run("未发现存取现与话单匹配的人员；")
            else:
                p.add_run("未发现存取现与话单匹配的人员；")
                
        except Exception as e:
            self.logger.warning(f"生成存取现与话单匹配人员信息时出错: {e}")
            p.add_run("未发现存取现与话单匹配的人员；")

    def _generate_large_fund_call_matching_persons(self, doc: Document, data_models: Dict):
        """生成大额资金跟踪与话单匹配的人员（性能优化版本）"""
        p = doc.add_paragraph()
        p.add_run("2.大额资金跟踪与话单匹配的人员：").bold = True
        
        try:
            # 使用预计算的大额交易数据
            large_amount_data = self._cached_data.get('large_amount_data', {})
            call_data = self._cached_data.get('call_data')
            
            if not large_amount_data or call_data is None:
                p.add_run("无大额交易数据或话单数据")
                return
            
            # 性能优化：批量处理大额资金跟踪
            fund_tracking_results = self._batch_fund_tracking(large_amount_data, call_data)
            
            if fund_tracking_results:
                p.add_run(f"共找到{len(fund_tracking_results)}个大额资金跟踪记录：")
                for i, (person, tracking_info) in enumerate(fund_tracking_results.items()):
                    if i < 10:  # 只显示前10个
                        p.add_run(f"{person}({tracking_info})")
                        if i < min(9, len(fund_tracking_results) - 1):
                            p.add_run("；")
                if len(fund_tracking_results) > 10:
                    p.add_run(f"等{len(fund_tracking_results)}人")
            else:
                p.add_run("无大额资金跟踪记录")
                
        except Exception as e:
            p.add_run(f"分析出错: {str(e)}")

    def _batch_fund_tracking(self, large_amount_data: Dict, call_data: pd.DataFrame) -> Dict:
        """批量处理大额资金跟踪，提升性能"""
        tracking_results = {}
        
        # 创建话单日期索引
        call_date_groups = call_data.groupby('date_key')
        
        for data_type, large_transactions in large_amount_data.items():
            # 批量处理每个数据源的大额交易
            for _, transaction in large_transactions.iterrows():
                person_name = transaction['本方姓名']
                tx_date = pd.to_datetime(transaction.get('交易日期', ''), errors='coerce')
                
                if pd.isna(tx_date):
                    continue
                
                date_key = tx_date.date()
                if date_key in call_date_groups.groups:
                    same_day_calls = call_date_groups.get_group(date_key)
                    
                    # 查找匹配的话单记录
                    person_calls = same_day_calls[
                        (same_day_calls['本方姓名'] == person_name) | 
                        (same_day_calls['对方姓名'] == person_name)
                    ]
                    
                    if not person_calls.empty:
                        # 收集联系人
                        contacts = set()
                        for _, call in person_calls.iterrows():
                            if call['本方姓名'] == person_name:
                                contacts.add(call['对方姓名'])
                            else:
                                contacts.add(call['本方姓名'])
                        
                        contacts.discard(person_name)  # 移除本人
                        if contacts:
                            if person_name not in tracking_results:
                                tracking_results[person_name] = []
                            tracking_results[person_name].extend(list(contacts))
        
        # 去重并格式化结果
        formatted_results = {}
        for person, contacts in tracking_results.items():
            unique_contacts = list(set(contacts))
            formatted_results[person] = f"{len(unique_contacts)}个联系人"
        
        return formatted_results

    def _generate_large_fund_call_matching_persons_original(self, doc: Document, data_models: Dict):
        """生成大额资金跟踪与话单匹配的人员（原始版本，保留作为备份）"""
        p = doc.add_paragraph()
        p.add_run("2.大额资金跟踪与话单匹配的人员：").bold = True
        
        # 使用大额资金追踪引擎
        fund_tracker = FundTrackingEngine(self.config)
        tracking_results = fund_tracker.track_large_funds(data_models)
        
        if not tracking_results.empty:
            try:
                # 获取与话单匹配的信息
                from src.export.excel_exporter import ExcelExporter
                excel_exporter = ExcelExporter()
                
                # 批量获取话单匹配信息
                person_names = tracking_results['核心人员'].tolist()
                transaction_dates = tracking_results['交易日期'].tolist()
                call_matches = excel_exporter._get_call_record_match_batch_optimized(
                    person_names, transaction_dates, data_models
                )
                
                # 将话单匹配信息添加到追踪结果中
                tracking_results['话单匹配'] = call_matches
                
                # 筛选出有话单匹配的记录
                matched_records = tracking_results[tracking_results['话单匹配'] != '']
                
                if not matched_records.empty:
                    # 收集匹配的人员信息
                    matched_persons = set()
                    for _, row in matched_records.iterrows():
                        person_name = row.get('核心人员', '未知')
                        match_info = row.get('话单匹配', '')
                        
                        # 从话单匹配信息中提取对方人员
                        if '对方' in match_info:
                            # 提取对方人员信息
                            parts = match_info.split('对方')
                            if len(parts) > 1:
                                opponents = parts[1].split(',')
                                for opponent in opponents:
                                    opponent = opponent.strip()
                                    if opponent:
                                        matched_persons.add(opponent)
                        
                        # 添加核心人员本身
                        if person_name and person_name != '未知':
                            matched_persons.add(person_name)
                    
                    # 完全列举所有匹配的人员信息
                    if matched_persons:
                        person_list = sorted(list(matched_persons))
                        person_with_units = []
                        for person in person_list:
                            unit_info = self._extract_unit_info_from_call_data(person, data_models)
                            if unit_info:
                                person_with_units.append(f"{person}（{unit_info}）")
                            else:
                                person_with_units.append(person)
                        p.add_run(f"发生大额资金当天与话单匹配的人员有{', '.join(person_with_units)}；")
                    else:
                        p.add_run("未发现大额资金与话单匹配的人员；")
                else:
                    p.add_run("未发现大额资金与话单匹配的人员；")
            except Exception as e:
                self.logger.warning(f"生成大额资金与话单匹配人员信息时出错: {e}")
                p.add_run("未发现大额资金与话单匹配的人员；")
        else:
            p.add_run("未发现大额资金与话单匹配的人员；")

    def _generate_large_fund_tracking_persons(self, doc: Document, data_models: Dict):
        """生成大额资金跟踪层级区分的重点人员（性能优化版本）"""
        p = doc.add_paragraph()
        p.add_run("3.大额资金跟踪层级区分的重点人员：").bold = True
        
        try:
            # 使用预计算的大额交易数据
            large_amount_data = self._cached_data.get('large_amount_data', {})
            
            if not large_amount_data:
                p.add_run("无大额交易数据")
                return
            
            # 性能优化：简化的层级分析
            level_results = self._batch_level_analysis(large_amount_data)
            
            if level_results:
                p.add_run(f"共找到{len(level_results)}个重点人员：")
                for i, (person, level_info) in enumerate(level_results.items()):
                    if i < 10:  # 只显示前10个
                        p.add_run(f"{person}({level_info})")
                        if i < min(9, len(level_results) - 1):
                            p.add_run("；")
                if len(level_results) > 10:
                    p.add_run(f"等{len(level_results)}人")
            else:
                p.add_run("无重点人员记录")
                
        except Exception as e:
            p.add_run(f"分析出错: {str(e)}")

    def _batch_level_analysis(self, large_amount_data: Dict) -> Dict:
        """批量处理层级分析，提升性能"""
        level_results = {}
        
        for data_type, large_transactions in large_amount_data.items():
            # 按人员分组统计
            person_stats = large_transactions.groupby('本方姓名').agg({
                '收入金额': ['sum', 'count'],
                '支出金额': ['sum', 'count']
            }).reset_index()
            
            # 扁平化列名
            person_stats.columns = ['person', 'income_sum', 'income_count', 'expense_sum', 'expense_count']
            
            for _, row in person_stats.iterrows():
                person = row['person']
                total_amount = row['income_sum'] + row['expense_sum']
                total_count = row['income_count'] + row['expense_count']
                
                # 简单的层级分类
                if total_amount > 1000000:  # 100万以上
                    level = "一级"
                elif total_amount > 500000:  # 50万以上
                    level = "二级"
                else:
                    level = "三级"
                
                level_results[person] = f"{level}({total_amount:,.0f}元,{total_count}笔)"
        
        return level_results

    def _generate_large_fund_tracking_persons_original(self, doc: Document, data_models: Dict):
        """生成大额资金跟踪层级区分的重点人员（原始版本，保留作为备份）"""
        p = doc.add_paragraph()
        p.add_run("3.大额资金跟踪层级区分的重点人员：").bold = True
        
        # 使用大额资金追踪引擎
        fund_tracker = FundTrackingEngine(self.config)
        tracking_results = fund_tracker.track_large_funds(data_models)
        
        if not tracking_results.empty:
            # 按追踪层级分组
            if '追踪层级' in tracking_results.columns and '核心人员' in tracking_results.columns:
                level_groups = tracking_results.groupby('追踪层级')
                
                level_info = []
                for level, group in level_groups:
                    persons = group['核心人员'].dropna().unique()
                    persons_list = list(set(persons) - {'未知', 'N/A', ''})  # 不限制每层显示人员数量
                    
                    if persons_list:
                        # 获取每个人的单位信息
                        person_with_units = []
                        for person in persons_list:
                            unit_info = self._extract_unit_info_from_call_data(person, data_models)
                            if unit_info:
                                person_with_units.append(f"{person}（{unit_info}）")
                            else:
                                person_with_units.append(f"{person}")
                        
                        level_info.append(f"跟踪层级为{level}的人员有{', '.join(person_with_units)}")
                
                if level_info:
                    p.add_run("；".join(level_info))
                    p.add_run("；")
                else:
                    p.add_run("未发现大额资金跟踪层级区分的重点人员；")
            else:
                p.add_run("未发现大额资金跟踪层级区分的重点人员；")
        else:
            p.add_run("未发现大额资金跟踪层级区分的重点人员；")

    def _extract_unit_info_from_call_data(self, person_name: str, data_models: Dict) -> str:
        """从话单数据中提取人员的单位信息"""
        try:
            if 'call' in data_models and data_models['call'] and not data_models['call'].data.empty:
                call_data = data_models['call'].data
                name_column = data_models['call'].name_column
                
                # 查找该人员的记录
                person_records = call_data[call_data[name_column] == person_name]
                
                if not person_records.empty:
                    # 尝试获取单位信息列
                    unit_columns = ['对方单位名称', '对方单位', '单位名称']
                    for col in unit_columns:
                        if col in person_records.columns:
                            units = person_records[col].dropna().unique()
                            if len(units) > 0:
                                # 返回第一个非空单位信息
                                return str(units[0])
            
            return ""
        except Exception as e:
            self.logger.warning(f"提取人员{person_name}的单位信息时出错: {e}")
            return ""

    def _save_document(self, doc: Document, title: str) -> Optional[str]:
        try:
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"{safe_title}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)
            self.logger.info(f"分析报告已导出到 {filepath}")
            return filepath
        except Exception as e:
            self.logger.error(f"保存Word报告时出错: {e}")
            return None