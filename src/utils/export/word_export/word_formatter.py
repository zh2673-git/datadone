"""
Word格式化器模块

提供Word文档格式化的相关功能，包括段落格式、表格格式等。
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Optional
import pandas as pd


class WordFormatter:
    """Word格式化器类"""
    
    def __init__(self):
        """初始化Word格式化器"""
        pass
    
    def add_df_to_doc(self, doc: Document, df: pd.DataFrame, title: Optional[str] = None) -> None:
        """
        将DataFrame添加到Word文档中
        
        Parameters:
        -----------
        doc : Document
            Word文档对象
        df : pd.DataFrame
            要添加的DataFrame
        title : Optional[str]
            表格标题
        """
        if df.empty:
            return
        
        if title:
            doc.add_paragraph(title)
        
        # 格式化数字列
        formatted_df = self._format_dataframe_numbers(df)
        
        # 添加表格到文档
        table = doc.add_table(rows=len(formatted_df) + 1, cols=len(formatted_df.columns))
        
        # 设置表头
        header_cells = table.rows[0].cells
        for i, col_name in enumerate(formatted_df.columns):
            header_cells[i].text = str(col_name)
            # 设置表头格式
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # 填充数据
        for i, (_, row) in enumerate(formatted_df.iterrows()):
            row_cells = table.rows[i + 1].cells
            for j, value in enumerate(row):
                row_cells[j].text = str(value)
        
        # 设置表格样式
        table.style = 'Table Grid'
        
        # 添加空行分隔
        doc.add_paragraph()
    
    def add_grouped_df_to_doc(self, doc: Document, df: pd.DataFrame, group_by_col: str, 
                             title: Optional[str] = None) -> None:
        """
        将分组DataFrame添加到Word文档中
        
        Parameters:
        -----------
        doc : Document
            Word文档对象
        df : pd.DataFrame
            要添加的DataFrame
        group_by_col : str
            分组列名
        title : Optional[str]
            表格标题
        """
        if df.empty or group_by_col not in df.columns:
            return
        
        if title:
            doc.add_heading(title, level=4)
        
        # 按分组列分组
        grouped = df.groupby(group_by_col)
        
        for group_name, group_df in grouped:
            # 添加分组标题
            doc.add_paragraph(f"{group_by_col}: {group_name}")
            
            # 添加分组数据表格
            self.add_df_to_doc(doc, group_df)
    
    def _format_dataframe_numbers(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        格式化DataFrame中的数字列
        
        Parameters:
        -----------
        df : pd.DataFrame
            要格式化的DataFrame
            
        Returns:
        --------
        pd.DataFrame
            格式化后的DataFrame
        """
        formatted_df = df.copy()
        
        for col in formatted_df.columns:
            # 检查是否为数值列
            if pd.api.types.is_numeric_dtype(formatted_df[col]):
                # 格式化数值列
                formatted_df[col] = formatted_df[col].apply(self._format_numeric_value)
        
        return formatted_df
    
    def _format_numeric_value(self, value) -> str:
        """
        格式化数值
        
        Parameters:
        -----------
        value : Any
            要格式化的值
            
        Returns:
        --------
        str
            格式化后的字符串
        """
        if pd.isna(value):
            return ""
        
        try:
            # 尝试转换为浮点数
            num_value = float(value)
            
            # 如果是整数，显示为整数格式
            if num_value == int(num_value):
                return f"{int(num_value):,}"
            else:
                # 浮点数保留两位小数
                return f"{num_value:,.2f}"
        except (ValueError, TypeError):
            # 如果无法转换为数字，返回原始字符串
            return str(value)
    
    def format_time_range_to_year_month(self, start_date, end_date) -> str:
        """
        格式化时间范围为年月格式
        
        Parameters:
        -----------
        start_date : datetime
            开始日期
        end_date : datetime
            结束日期
            
        Returns:
        --------
        str
            格式化后的时间范围字符串
        """
        if pd.isna(start_date) or pd.isna(end_date):
            return "未知时间范围"
        
        try:
            start_str = start_date.strftime('%Y年%m月')
            end_str = end_date.strftime('%Y年%m月')
            return f"{start_str}至{end_str}"
        except AttributeError:
            return "未知时间范围"
    
    def is_numeric_value(self, value) -> bool:
        """
        检查值是否为数值
        
        Parameters:
        -----------
        value : Any
            要检查的值
            
        Returns:
        --------
        bool
            是否为数值
        """
        if pd.isna(value):
            return False
        
        try:
            float(value)
            return True
        except (ValueError, TypeError):
            return False
    
    def to_chinese_numeral(self, num: int) -> str:
        """
        将数字转换为中文数字
        
        Parameters:
        -----------
        num : int
            要转换的数字
            
        Returns:
        --------
        str
            中文数字
        """
        chinese_numerals = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        
        if num <= 0:
            return "零"
        elif num <= 10:
            return chinese_numerals[num]
        elif num < 20:
            return f"十{chinese_numerals[num - 10]}"
        elif num < 100:
            tens = num // 10
            units = num % 10
            if units == 0:
                return f"{chinese_numerals[tens]}十"
            else:
                return f"{chinese_numerals[tens]}十{chinese_numerals[units]}"
        else:
            return str(num)
    
    def add_centered_heading(self, doc: Document, text: str, level: int = 1) -> None:
        """
        添加居中的标题
        
        Parameters:
        -----------
        doc : Document
            Word文档对象
        text : str
            标题文本
        level : int
            标题级别
        """
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_formatted_paragraph(self, doc: Document, text: str, bold: bool = False, 
                               italic: bool = False, font_size: Optional[int] = None) -> None:
        """
        添加格式化的段落
        
        Parameters:
        -----------
        doc : Document
            Word文档对象
        text : str
            段落文本
        bold : bool
            是否加粗
        italic : bool
            是否斜体
        font_size : Optional[int]
            字体大小
        """
        paragraph = doc.add_paragraph(text)
        
        if bold or italic or font_size:
            for run in paragraph.runs:
                if bold:
                    run.font.bold = True
                if italic:
                    run.font.italic = True
                if font_size:
                    run.font.size = Pt(font_size)
    
    def add_bullet_list(self, doc: Document, items: List[str]) -> None:
        """
        添加项目符号列表
        
        Parameters:
        -----------
        doc : Document
            Word文档对象
        items : List[str]
            项目列表
        """
        for item in items:
            paragraph = doc.add_paragraph(item)
            paragraph.style = 'List Bullet'
    
    def add_numbered_list(self, doc: Document, items: List[str]) -> None:
        """
        添加编号列表
        
        Parameters:
        -----------
        doc : Document
            Word文档对象
        items : List[str]
            项目列表
        """
        for i, item in enumerate(items, 1):
            paragraph = doc.add_paragraph(f"{i}. {item}")
            paragraph.style = 'List Number'