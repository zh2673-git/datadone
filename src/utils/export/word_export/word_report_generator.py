"""
Word报告生成器模块

提供Word报告生成的相关功能，包括个人分析报告、综合报告等。
"""

import pandas as pd
from typing import Dict, List, Optional, Any
from docx import Document
import logging

from .word_formatter import WordFormatter


class WordReportGenerator:
    """Word报告生成器类"""
    
    def __init__(self):
        """初始化Word报告生成器"""
        self.logger = logging.getLogger(self.__class__.__name__)
        self.formatter = WordFormatter()
    
    def generate_person_bank_analysis(self, doc: Document, person_name: str, analyzer) -> bool:
        """
        生成个人银行分析报告
        
        Parameters:
        -----------
        doc : Document
            Word文档对象
        person_name : str
            个人姓名
        analyzer
            银行分析器对象
            
        Returns:
        --------
        bool
            是否成功生成报告
        """
        try:
            person_data = analyzer.bank_model.get_data_by_person(person_name)
            if person_data.empty:
                doc.add_paragraph(f"未找到 {person_name} 的银行数据。")
                return False

            section_num = 1
            has_content = False

            # 转账分析
            transfer_freq_df = analyzer.analyze_frequency(person_data)
            if not transfer_freq_df.empty:
                self._generate_frequency_summary_paragraph(doc, person_name, transfer_freq_df, person_data, "银行转账", analyzer.bank_model, section_num)
                section_num += 1
                self._add_top_opponent_tables(doc, transfer_freq_df)
                has_content = True
            
            # 存取现分析
            cash_ops_df = analyzer.analyze_cash_operations(person_data)
            if not cash_ops_df.empty:
                deposit_df = cash_ops_df[cash_ops_df['存取现标识'] == '存现']
                if not deposit_df.empty:
                    self._generate_cash_summary_paragraph(doc, person_name, deposit_df, '存现', analyzer.bank_model, section_num)
                    section_num += 1
                    has_content = True
                
                withdraw_df = cash_ops_df[cash_ops_df['存取现标识'] == '取现']
                if not withdraw_df.empty:
                    self._generate_cash_summary_paragraph(doc, person_name, withdraw_df, '取现', analyzer.bank_model, section_num)
                    section_num += 1
                    has_content = True
            
            # 特殊分析
            special_amounts_df = analyzer.analyze_special_amounts(person_data.copy())
            special_dates_df = analyzer.analyze_special_dates(person_data.copy())
            if not special_amounts_df.empty or not special_dates_df.empty:
                self._generate_special_analysis_summary(doc, person_name, analyzer, special_amounts_df, special_dates_df, "银行", section_num)
                section_num += 1
                has_content = True

            # 整数金额分析
            integer_amounts_df = analyzer.analyze_integer_amounts(person_data.copy())
            if not integer_amounts_df.empty:
                self._generate_integer_amount_summary(doc, person_name, analyzer, integer_amounts_df, "银行", section_num)
                section_num += 1
                has_content = True

            # 重点收支分析
            self._generate_key_transactions_summary(doc, person_name, analyzer, person_data, section_num)
            has_content = True
            
            if not has_content:
                doc.add_paragraph(f"未找到 {person_name} 的有效银行交易数据。")
            
            return has_content
            
        except Exception as e:
            self.logger.error(f"生成个人银行分析报告时出错: {e}")
            return False
    
    def generate_person_payment_analysis(self, doc: Document, person_name: str, analyzer, payment_type: str) -> bool:
        """
        生成个人支付平台分析报告
        
        Parameters:
        -----------
        doc : Document
            Word文档对象
        person_name : str
            个人姓名
        analyzer
            支付平台分析器对象
        payment_type : str
            支付平台类型（微信/支付宝）
            
        Returns:
        --------
        bool
            是否成功生成报告
        """
        try:
            person_data = analyzer.data_model.get_data_by_person(person_name)
            if person_data.empty:
                doc.add_paragraph(f"未找到 {person_name} 的{payment_type}数据。")
                return False
            
            section_num = 1
            has_content = False

            freq_df = analyzer.analyze_frequency(person_data)
            if not freq_df.empty:
                self._generate_frequency_summary_paragraph(doc, person_name, freq_df, person_data, f"{payment_type}交易", analyzer.data_model, section_num)
                section_num += 1
                self._add_top_opponent_tables(doc, freq_df)
                has_content = True
            
            # 特殊分析
            special_amounts_df = analyzer.analyze_special_amounts(person_data.copy())
            special_dates_df = analyzer.analyze_special_dates(person_data.copy())
            if not special_amounts_df.empty or not special_dates_df.empty:
                self._generate_special_analysis_summary(doc, person_name, analyzer, special_amounts_df, special_dates_df, payment_type, section_num)
                section_num += 1
                has_content = True

            # 整数金额分析
            platform_type = 'wechat' if payment_type == '微信' else 'alipay'
            integer_amounts_df = analyzer.analyze_integer_amounts(person_data.copy(), platform_type)
            if not integer_amounts_df.empty:
                self._generate_integer_amount_summary(doc, person_name, analyzer, integer_amounts_df, payment_type, section_num)
                section_num += 1
                has_content = True

            # 重点收支分析
            self._generate_payment_key_transactions_summary(doc, person_name, analyzer, payment_type, section_num)
            has_content = True

            if not has_content:
                doc.add_paragraph(f"未找到 {person_name} 的{payment_type}交易数据。")
            
            return has_content
            
        except Exception as e:
            self.logger.error(f"生成个人{payment_type}分析报告时出错: {e}")
            return False
    
    def generate_person_call_analysis(self, doc: Document, person_name: str, analyzer) -> bool:
        """
        生成个人通话分析报告
        
        Parameters:
        -----------
        doc : Document
            Word文档对象
        person_name : str
            个人姓名
        analyzer
            通话分析器对象
            
        Returns:
        --------
        bool
            是否成功生成报告
        """
        try:
            person_data = analyzer.call_model.get_data_by_person(person_name)
            if person_data.empty:
                doc.add_paragraph(f"未找到 {person_name} 的话单数据。")
                return False

            freq_df = analyzer.analyze_call_frequency(person_data)
            if not freq_df.empty:
                self._generate_call_summary_paragraph(doc, person_name, freq_df, person_data)
                
                # 过滤掉对方姓名为"未知"的记录
                known_contacts = freq_df[~freq_df['对方姓名'].isin(['未知', 'N/A', '']) & ~freq_df['对方姓名'].isna()]
                
                # 如果没有已知联系人，尝试使用其他数据
                if known_contacts.empty:
                    # 使用所有数据，但后续会按号码等其他信息排序
                    known_contacts = freq_df.copy()
                    # 按通话次数排序
                    top_contacts = known_contacts.nlargest(5, '通话次数')
                else:
                    # 只使用已知联系人，按通话次数排序
                    top_contacts = known_contacts.nlargest(5, '通话次数')
                
                # 如果有数据可显示
                if not top_contacts.empty:
                    doc.add_paragraph(f"最密切联系人TOP {min(5, len(top_contacts))}:")
                    
                    # 创建一个新的DataFrame用于显示
                    display_df = pd.DataFrame()
                    
                    # 复制基本列
                    display_df['对方姓名'] = top_contacts['对方姓名']
                    display_df['对方号码'] = top_contacts['对方号码']
                    display_df['通话次数'] = top_contacts['通话次数']
                    display_df['通话总时长(分钟)'] = top_contacts['通话总时长(分钟)']
                    
                    # 处理对方单位列
                    # 首先检查是否有带后缀的对方单位列（来自话单数据）
                    unit_col = next((col for col in top_contacts.columns if '对方单位名称_' in col), None)
                    if unit_col:
                        display_df['对方单位'] = top_contacts[unit_col]
                    elif '对方单位' in top_contacts.columns:
                        # 如果有多个单位（用|分隔），只取第一个
                        display_df['对方单位'] = top_contacts['对方单位'].apply(
                            lambda x: x.split('|')[0] if pd.notna(x) and '|' in str(x) else x
                        )
                    else:
                        display_df['对方单位'] = None
                    
                    # 填充空值
                    display_df['对方单位'] = display_df['对方单位'].fillna('N/A')
                    
                    # 定义最终显示列的顺序
                    display_columns = [
                        '对方姓名', '对方单位', '对方号码', 
                        '通话次数', '通话总时长(分钟)'
                    ]
                    
                    # 添加到文档
                    self.formatter.add_df_to_doc(doc, display_df[display_columns])
                else:
                    doc.add_paragraph("没有找到有效的联系人信息。")
                
                return True
            else:
                doc.add_paragraph(f"未找到 {person_name} 的通话频率数据。")
                return False
            
        except Exception as e:
            self.logger.error(f"生成个人通话分析报告时出错: {e}")
            return False
    
    def _generate_frequency_summary_paragraph(self, doc: Document, person_name: str, frequency_df: pd.DataFrame, 
                                             person_data: pd.DataFrame, analysis_type_str: str, data_model, section_num: int) -> None:
        """生成频率分析概要段落"""
        total_income = frequency_df['总收入'].sum()
        total_expense = frequency_df['总支出'].abs().sum()
        net_flow = total_income - total_expense
        total_transactions_count = frequency_df['交易次数'].sum()

        start_date = person_data[data_model.date_column].min()
        end_date = person_data[data_model.date_column].max()
        
        # 检查是否为NaT值，如果是则使用默认字符串
        if pd.isna(start_date):
            start_date_str = "未知开始日期"
        else:
            start_date_str = start_date.strftime('%Y-%m-%d')
        
        if pd.isna(end_date):
            end_date_str = "未知结束日期"
        else:
            end_date_str = end_date.strftime('%Y-%m-%d')
        
        time_span_str = f"{start_date_str}至{end_date_str}"

        # 计算主要时间集中
        person_data = person_data.copy()
        person_data['年份月份'] = person_data[data_model.date_column].dt.strftime('%Y年%m月')
        monthly_counts = person_data.groupby('年份月份').size().reset_index(name='次数')
        top_months = monthly_counts.nlargest(3, '次数')
        major_time_str = ", ".join([f"{row['年份月份']} ({row['次数']}次)" for _, row in top_months.iterrows()])

        # 单笔主要金额，添加对方姓名，避免重复对象
        all_amounts = person_data.sort_values(by=data_model.amount_column, ascending=False)
        top_amounts_with_names = []

        # 跟踪已添加的对象，避免重复
        added_opponents = set()

        # 获取前三笔金额，优先选择不同对手方
        for _, row in all_amounts.iterrows():
            amount = row[data_model.amount_column]
            opponent = row[data_model.opposite_name_column]
            opponent_str = str(opponent) if pd.notna(opponent) else "未知"

            # 如果这个对手方还没有被添加过，或者已经有3个不同的对手方了
            if opponent_str not in added_opponents and len(top_amounts_with_names) < 3:
                formatted_amount = f"{amount:.2f}"
                top_amounts_with_names.append(f"{formatted_amount}元（{opponent_str}）")
                added_opponents.add(opponent_str)
        
        top_single_amounts = "、".join(top_amounts_with_names)

        # 重复最多的金额（前三名）
        amount_counts = person_data[data_model.amount_column].value_counts()
        top_frequent_amounts = []
        for i in range(min(3, len(amount_counts))):
            amount = amount_counts.index[i]
            count = amount_counts.iloc[i]
            top_frequent_amounts.append(f"{amount:.2f}元 ({count}次)")
        most_frequent_amount_info = "、".join(top_frequent_amounts)

        # 生成概览段落
        p = doc.add_paragraph()
        p.add_run(f"{section_num}、{analysis_type_str}概览").bold = True
        p.add_run(f"：在 {time_span_str} 期间，总收入 {total_income:.2f} 元，总支出 {total_expense:.2f} 元，")
        p.add_run(f"净流水 {net_flow:.2f} 元，共计 {total_transactions_count} 笔交易。")
        p.add_run(f"主要时间集中在：{major_time_str}。")
        p.add_run(f'单笔主要金额：{top_single_amounts}。')

        # 检查是否有大额金额需要加粗显示
        has_large_amount = any(abs(amount_counts.index[i]) > 5000 for i in range(min(3, len(amount_counts))))

        if has_large_amount:
            p.add_run('重复最多的金额为 ')
            # 对大额金额加粗显示
            for i, amount_info in enumerate(top_frequent_amounts):
                amount = amount_counts.index[i]
                if abs(amount) > 5000:
                    p.add_run(amount_info).bold = True
                else:
                    p.add_run(amount_info)
                if i < len(top_frequent_amounts) - 1:
                    p.add_run('、')
            p.add_run('。')
        else:
            p.add_run(f'重复最多的金额为：{most_frequent_amount_info}。')
    
    def _generate_call_summary_paragraph(self, doc: Document, person_name: str, freq_df: pd.DataFrame, person_data: pd.DataFrame) -> None:
        """生成话单分析的概要段落"""
        total_contacts = freq_df['对方号码'].nunique()
        total_calls = freq_df['通话次数'].sum()
        total_duration_min = freq_df['通话总时长(分钟)'].sum()
        
        start_date = person_data[person_data.columns[person_data.columns.str.contains('日期')][0]].min()
        end_date = person_data[person_data.columns[person_data.columns.str.contains('日期')][0]].max()
        
        # 检查是否为NaT值，如果是则使用默认字符串
        if pd.isna(start_date):
            start_date_str = "未知开始日期"
        else:
            start_date_str = start_date.strftime('%Y-%m-%d')
        
        if pd.isna(end_date):
            end_date_str = "未知结束日期"
        else:
            end_date_str = end_date.strftime('%Y-%m-%d')
        
        time_span_str = f"{start_date_str}至{end_date_str}"
        
        main_outgoing = person_data.get('主叫次数', pd.Series(0)).sum()
        main_incoming = person_data.get('被叫次数', pd.Series(0)).sum()

        summary = (
            f"在 {time_span_str} 期间，{person_name} 共通话 {total_calls} 次，总时长 {total_duration_min:,.2f} 分钟，联系了 {total_contacts} 个不同对象。"
            f"其中主叫 {main_outgoing} 次，被叫 {main_incoming} 次。"
        )
        doc.add_paragraph(summary)
    
    def _add_top_opponent_tables(self, doc: Document, freq_df: pd.DataFrame) -> None:
        """添加前几名对手方表格"""
        # 按收入排序的前几名
        top_income = freq_df.nlargest(5, '总收入')
        if not top_income.empty:
            doc.add_paragraph("收入最多的前几名对手方：")
            self.formatter.add_df_to_doc(doc, top_income[['对方姓名', '总收入', '交易次数']])
        
        # 按支出排序的前几名
        top_expense = freq_df.nlargest(5, '总支出')
        if not top_expense.empty:
            doc.add_paragraph("支出最多的前几名对手方：")
            self.formatter.add_df_to_doc(doc, top_expense[['对方姓名', '总支出', '交易次数']])
    
    def _generate_cash_summary_paragraph(self, doc: Document, person_name: str, cash_df: pd.DataFrame, 
                                        cash_type: str, data_model, section_num: int) -> None:
        """生成存取现分析概要段落"""
        total_amount = cash_df['交易金额'].abs().sum()
        transaction_count = len(cash_df)
        
        # 计算平均金额
        avg_amount = total_amount / transaction_count if transaction_count > 0 else 0
        
        # 获取时间范围
        start_date = cash_df[data_model.date_column].min()
        end_date = cash_df[data_model.date_column].max()
        
        if pd.isna(start_date):
            start_date_str = "未知开始日期"
        else:
            start_date_str = start_date.strftime('%Y-%m-%d')
        
        if pd.isna(end_date):
            end_date_str = "未知结束日期"
        else:
            end_date_str = end_date.strftime('%Y-%m-%d')
        
        time_span_str = f"{start_date_str}至{end_date_str}"
        
        summary = (
            f"{section_num}、{cash_type}分析：在 {time_span_str} 期间，共发生 {transaction_count} 笔{cash_type}交易，"
            f"总金额 {total_amount:.2f} 元，平均每笔 {avg_amount:.2f} 元。"
        )
        doc.add_paragraph(summary)
    
    def _generate_special_analysis_summary(self, doc: Document, person_name: str, analyzer, 
                                          special_amounts_df: pd.DataFrame, special_dates_df: pd.DataFrame, 
                                          analysis_type: str, section_num: int) -> None:
        """生成特殊分析概要"""
        doc.add_paragraph(f"{section_num}、{analysis_type}特殊分析：")
        
        if not special_amounts_df.empty:
            doc.add_paragraph("特殊金额分析：")
            self.formatter.add_df_to_doc(doc, special_amounts_df)
        
        if not special_dates_df.empty:
            doc.add_paragraph("特殊日期分析：")
            self.formatter.add_df_to_doc(doc, special_dates_df)
    
    def _generate_integer_amount_summary(self, doc: Document, person_name: str, analyzer, 
                                        integer_amounts_df: pd.DataFrame, analysis_type: str, section_num: int) -> None:
        """生成整数金额分析概要"""
        doc.add_paragraph(f"{section_num}、{analysis_type}整数金额分析：")
        self.formatter.add_df_to_doc(doc, integer_amounts_df)
    
    def _generate_key_transactions_summary(self, doc: Document, person_name: str, analyzer, 
                                          person_data: pd.DataFrame, section_num: int) -> None:
        """生成重点收支分析概要"""
        try:
            key_transactions = analyzer.analyze_key_transactions(person_data)
            if not key_transactions.empty:
                doc.add_paragraph(f"{section_num}、重点收支分析：")
                self.formatter.add_df_to_doc(doc, key_transactions)
        except Exception as e:
            self.logger.warning(f"重点收支分析失败: {e}")
    
    def _generate_payment_key_transactions_summary(self, doc: Document, person_name: str, analyzer, 
                                                  payment_type: str, section_num: int) -> None:
        """生成支付平台重点收支分析概要"""
        try:
            key_transactions = analyzer.analyze_key_transactions()
            if not key_transactions.empty:
                doc.add_paragraph(f"{section_num}、{payment_type}重点收支分析：")
                self.formatter.add_df_to_doc(doc, key_transactions)
        except Exception as e:
            self.logger.warning(f"{payment_type}重点收支分析失败: {e}")